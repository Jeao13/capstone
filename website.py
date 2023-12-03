from flask import jsonify
from flask import Flask, render_template, request, redirect, url_for, session, flash, send_file, jsonify, make_response, Response
import mysql.connector
from mysql.connector import pooling
import base64
import io
import os
import requests
import convertapi
import detectlanguage
import random
import string
import base64
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from werkzeug.utils import secure_filename
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import letter
from datetime import datetime
import pandas as pd
import numpy as np
from scipy.stats import randint
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.model_selection import train_test_split
from sklearn.svm import LinearSVC
from sklearn.model_selection import cross_val_score
import time
from docx.oxml import OxmlElement
from PIL import Image
import io
from docx.shared import Inches
import subprocess





def create_connection_pool():
    db_config = {
    'host': os.environ.get('MYSQL_HOST', 'mysql-uetk'),
    'user': os.environ.get('MYSQL_USER', 'mysql'),
    'password': os.environ.get('MYSQL_PASSWORD', '1NYNmyNJSq59o8UBx3d57qFZehQyl/GfjICwd6/PpgE='),
    'database': os.environ.get('MYSQL_DATABASE', 'mysql'),
    'port': os.environ.get('MYSQL_PORT', '3306'),
    }
        
    cnxpool = pooling.MySQLConnectionPool(pool_name = "example_pool", pool_size = 20, autocommit=True,  **db_config)

    return cnxpool







app = Flask(__name__)
app.secret_key = 'your_secret_key'  # Set a secret key for session management





pdfkit_options = {
    'page-size': 'Custom', 
    'page-width': '215.9mm',  # 8.5 inches converted to millimeters
    'page-height': '330.2mm',  # 13 inches converted to millimeters
    'encoding': 'UTF-8',
}



@app.route('/get_data_endpoint', methods=['GET'])
def get_data(x):
    data = x
    return jsonify(data)




def notifs(user_id, message):

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()
    db_cursor.execute("INSERT INTO notifications (user_id  ,message) VALUES (%s, %s)",
                      (user_id, message))
    cursor1.commit()

    db_cursor.close()


def replace_placeholder1(doc, placeholder, image_path, font_size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT, bold=False, indentation_spaces=0):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            for run in paragraph.runs:
                run.clear()

            # Add the image to the paragraph
            indentation = "               " * indentation_spaces
            paragraph.alignment = alignment
            paragraph.left_indent = Cm(10)
            run = paragraph.add_run()
            run.bold = bold
            run.font.size = Pt(font_size)

            run.add_text(indentation)
            # Adjust width and height as needed
            run.add_picture(image_path, width=Cm(2.88), height=Cm(1.56))


def replace_placeholder(doc, placeholder, new_text, font_size=12, bold=False, italic=False, alignment=None):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, new_text)
                    run.font.size = Pt(font_size)
                    run.font.bold = bold
                    run.font.italic = italic
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
                    if alignment:
                        run.alignment = alignment


def clear_and_add_line(doc, line_number, new_text, font_size=12, bold=False, italic=False, alignment=None, indentation=None):
    for i, paragraph in enumerate(doc.paragraphs):
        if i == line_number:
            for run in paragraph.runs:
                run.clear()
            run = paragraph.add_run(new_text)
            run.font.size = Pt(font_size)
            run.font.bold = bold
            run.font.italic = italic
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
            if alignment:
                run.alignment = alignment

            if indentation is not None:
                # Set the left indentation for the paragraph
                paragraph.paragraph_format.left_indent = Pt(indentation)

            break


def replace_table_cell_placeholder(table, row_index, col_index, new_text):
    cell = table.cell(row_index, col_index)

    # Clear the existing text in the cell
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""

    # Add the new text to the cell
    cell.text = new_text


def replace_table_cell_placeholder1(table, row_index, col_index, new_text, placeholder):
    cell = table.cell(row_index, col_index)

    # Flag to track if the placeholder was found and replaced
    placeholder_replaced = False
    if new_text == "" or None:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if placeholder in run.text:

                    run.text = run.text.replace(placeholder, "N/A")
                    placeholder_replaced = True

        # Add the new text to the cell if the placeholder was not found
        if not placeholder_replaced:
            cell.text = new_text
    else:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if placeholder in run.text:

                    run.text = run.text.replace(placeholder, new_text)
                    placeholder_replaced = True

        # Add the new text to the cell if the placeholder was not found
        if not placeholder_replaced:
            cell.text = new_text


def replace_table_cell_placeholder2(table, row_index, col_index, new_text, placeholder):
    cell = table.cell(row_index, col_index)

    if new_text == "checked":
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if placeholder in run.text:

                    run.text = run.text.replace(placeholder, "☑")

    else:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if placeholder in run.text:

                    run.text = run.text.replace(placeholder, "☐")
                    pdfkit_options


def toggle_table_cell_checkbox(table, row_index, col_index, status):
    cell = table.cell(row_index, col_index)
    print(status)
    if status == "checked":
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.text = ""
        cell.text = "☑"

        # Center-align the text horizontally
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)  # Adjust font size if needed
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    else:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.text = ""
        cell.text = "☐"

        # Center-align the text horizontally
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)  # Adjust font size if needed
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def replace_table_cell_placeholder_with_image(table, row_index, col_index, image_path, placeholder, indentation_spaces=0):
    cell = table.cell(row_index, col_index)

    # Iterate through the runs in the cell
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if placeholder in run.text:
                # Clear the existing run
                run.clear()

                # Add an image with the specified path
                indentation = "   " * indentation_spaces
                paragraph.left_indent = Cm(3)
                run = paragraph.add_run()
                run.add_text(indentation)
                run.add_picture(image_path, width=Cm(2.88), height=Cm(1.56))


def generate_random_code(length=8):
    # Define the characters to choose from
    characters = string.ascii_uppercase + string.digits

    # Generate a random code
    code = '#' + ''.join(random.choice(characters) for _ in range(length - 1))

    return code


@app.route('/submit_notice', methods=['GET', 'POST'])
def submit_notice():

    id = request.form.get('id')
    print(id)
    code = request.form.get('code')
    student = request.form.get('student')
    complainant = request.form.get('complainant')
    srcode = request.form.get('srcode')
    section = request.form.get('section')
    program = request.form.get('department')
    gender = request.form.get('gender')
    offense = request.form.get('offense_type')
    offense1 = request.form.get('offense_type1')
    minor_input1 = request.form.get('sanctionsminor')
    minor_input2 = request.form.get('sanctionsminor1')
    major_input1 = request.form.get('sanctionsmajor')
    major_input2 = request.form.get('sanctionsmajor1')
    fieldwork1 = request.form.get('fieldwork')
    prolonged1 = request.form.get('prolonged')
    fieldwork2 = request.form.get('fieldwork1')
    prolonged2 = request.form.get('prolonged1')
    specify3 = request.form.get('specify2')
    specify4 = request.form.get('specify3')
    print(specify3)
    print(specify4)
    statusreport = request.form.get('status')
    current_datetime = datetime.now()
    random_code = generate_random_code()
    current_date = current_datetime.date()
    formatted_date = current_date.strftime("%m/%d/%Y")
    current_time = datetime.now()


    specify2 = specify3 or specify4

    if specify2 is None:
        specify2 = "N/A"

    minor_input = minor_input1 or minor_input2
    major_input = major_input1 or major_input2

    if minor_input:
        major_input = ""

    else:
        minor_input = ""

# Format the current time as "hh:mm AM/PM"
    formatted_time = current_time.strftime("%I:%M %p")


    if program == "CAFAD":
            Name_Coordinator = "Paula Joyce A. Buisan"

    elif program == "CICS":
        Name_Coordinator = "Lovely Rose Tipan Hernandez"


    elif program == "CIT":
        Name_Coordinator = "Dolfus G. Miciano"


    elif program == "COE":
        Name_Coordinator = "Lovely Rose Tipan Hernandez"

    if gender == "male":
        status = "checked"
        status1 = "not"
    else:
        status1 = "checked"
        status = "not"


    print(offense)

    if offense == "minor":
        status3 = "checked"
        status2 = "not"
    elif offense == "major":
        print("wow")
        status2 = "checked"
        status3 = "not"

    if offense1 == "minor":
        status3 = "checked"
        status2 = "not"
    elif offense1 == "major":
        status2 = "checked"
        status3 = "not"

    if fieldwork1 == "fieldwork":
        status4 = "checked"
        
    else:
        status4 = "not"


    if fieldwork2 == "fieldwork1":
        status4 = "checked"
        
    else:
        status4 = "not"

    if prolonged1 == "prolonged":
        status5 = "checked"
        
    else:
        status5 = "not"


    if prolonged2 == "prolonged1":
        status5 = "checked"
        
    else:
        status5 = "not"
        

    pdf_filename = 'notice.docx'
    doc = Document(pdf_filename)

    toggle_table_cell_checkbox(doc.tables[0], 4, 19, status1)
    toggle_table_cell_checkbox(doc.tables[0], 4, 14, status)
    toggle_table_cell_checkbox(doc.tables[0], 8, 0, status3)
    toggle_table_cell_checkbox(doc.tables[0], 8, 8, status2)
    toggle_table_cell_checkbox(doc.tables[0], 10, 0, status4)
    toggle_table_cell_checkbox(doc.tables[0], 11, 0, status5)

    replace_table_cell_placeholder1(doc.tables[0], 2, 6, formatted_date, "(date)")
    replace_table_cell_placeholder1(doc.tables[0], 14, 2, formatted_date, "(date2)")
    replace_table_cell_placeholder1(doc.tables[0], 5, 6, program, "(program)")
    replace_table_cell_placeholder1(doc.tables[0], 3, 6, student, "(name)")
    replace_table_cell_placeholder1(doc.tables[0], 3, 18, srcode, "(code)")
    replace_table_cell_placeholder1(doc.tables[0], 5, 18, section, "(section)")
    replace_table_cell_placeholder1(doc.tables[0], 7, 6, minor_input, "(minor)")
    replace_table_cell_placeholder1(doc.tables[0], 7, 13, major_input, "(major)")
    replace_table_cell_placeholder1(doc.tables[0], 11, 12, specify2, "(specify)")
    replace_table_cell_placeholder1(doc.tables[0], 14, 2, Name_Coordinator, "NAME")

    toggle_table_cell_checkbox(doc.tables[1], 4, 19, status1)
    toggle_table_cell_checkbox(doc.tables[1], 4, 14, status)
    toggle_table_cell_checkbox(doc.tables[1], 8, 0, status3)
    toggle_table_cell_checkbox(doc.tables[1], 8, 8, status2)
    toggle_table_cell_checkbox(doc.tables[1], 10, 0, status4)
    toggle_table_cell_checkbox(doc.tables[1], 11, 0, status5)

    replace_table_cell_placeholder1(doc.tables[1], 2, 6, formatted_date, "(date)")
    replace_table_cell_placeholder1(doc.tables[1], 14, 2, formatted_date, "(date2)")
    replace_table_cell_placeholder1(doc.tables[1], 5, 6, program, "(program)")
    replace_table_cell_placeholder1(doc.tables[1], 3, 6, student, "(name)")
    replace_table_cell_placeholder1(doc.tables[1], 3, 18, srcode, "(code)")
    replace_table_cell_placeholder1(doc.tables[1], 5, 18, section, "(section)")
    replace_table_cell_placeholder1(doc.tables[1], 7, 6, minor_input, "(minor)")
    replace_table_cell_placeholder1(doc.tables[1], 7, 13, major_input, "(major)")
    replace_table_cell_placeholder1(doc.tables[1], 11, 12, specify2, "(specify)")
    replace_table_cell_placeholder1(doc.tables[1], 14, 2, Name_Coordinator, "NAME")

    doc.save("modified_document.docx")


    pdfpath = os.path.join('modified_document.docx')
   

    file_name = f'{random_code}_Notice of Case Dismissal'
    with open(pdfpath, "rb") as pdf_file:
        pdf_data = pdf_file.read()

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()
    db_cursor.execute("INSERT INTO notice_case (notice_id  ,complainant, name, coord, date, time, file, file_name,status) VALUES (%s, %s,%s, %s, %s, %s, %s, %s,%s)",
                      (random_code, complainant, student, Name_Coordinator, current_date, formatted_time, pdf_data, file_name, statusreport))
    cursor1.commit()

    db_cursor.close()

   

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor_status = cursor1.cursor()
    db_cursor_status.execute("UPDATE reports SET status = %s WHERE report_id = %s", (statusreport, code))
    cursor1.commit()
    db_cursor_status.close()

    flash('The report is submitted', 'success')
    return redirect('/head')


@app.route('/generate_report', methods=['GET', 'POST'])
def generate_report():

    form = request.form.get('form')
    to = request.form.get('to')

    print(form)
    print(to)

    minor = ["12.1.1 - attendance, punctuality, cutting classes", 
             "12.1.2 - dress code, uniform", "12.1.3 - property misuse", 
             "12.1.4 - noise disturbance", "12.1.5 - posting violation",
             "12.1.6 - notice removal","12.1.7 - littering",
             "12.1.8 - smoking violation","12.1.9 - trespassing",
             "12.1.10 - misconduct","12.1.11 - harassment",
             "12.1.12 - provocation, fight","12.1.13 - PDA",
             "12.1.14 - truancy",
             ]
    
    major = ["13.1 - repeat offenses","13.2 - insubordination","13.3 - smoking violation",
             "13.4 - alcohol violation","13.5 - intoxication","13.6 - trespassing",
             "13.7 - property misuse","13.8 - Reckless endangerment","13.9 - Gambling","13.10 - Identity fraud",
             "13.11 - Misuse of university name/logo","13.12 - Unauthorized representation","13.13 - abusive behavior","13.14 - unauthorized membership",
             "13.15 - online misconduct","13.16 - vandalism","13.17 - academic disruption",
             "13.18 - solicitation","13.19 - physical harm","13.20 - weapons possession",
             "13.21 - theft","13.22 - bribery","13.23 - sexual misconduct","13.24 - obscenity",
             "13.25 - defamation","13.26 - physical harm","13.27 - falsification","13.28 - disrepute",
             "13.29 - riot","13.30 - destruction of property","13.31 - burglary","13.32 - hazing",
             "13.33 - drugs","13.34 - firearms possession","13.35 - threats","13.36 - felonies","13.37 - moral turpitude",
             "14.1 - cheating, mobile phone","14.2 - cheating, talking","14.3 - cheating, dictating answers",
             "14.4 - cheating, notes possession","14.5 - cheating, outside information","14.6 - cheating, leakage facilitation",
             "14.7 - cheating, buying/selling questions","14.8 - cheating, copying answers","14.9 - cheating, covert devices",
             "14.10 - cheating, impersonation","14.11 - plagiarism","14.12 - cheating, surrogate attendance",
             "14.13 - plagiarism","14.14 - cheating, caught","14.15 - cheating, aiding"
             ]

    current_datetime = datetime.now()
    random_code = generate_random_code()
    current_date = current_datetime.date()
    formatted_date = current_date.strftime("%m/%d/%Y")
    current_time = datetime.now()

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()
    db_cursor.execute(
        "SELECT COUNT(*) FROM reports WHERE DATE(date_time) >= %s AND DATE(date_time) <= %s;", (form, to))
    result = db_cursor.fetchone()
    db_cursor.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor9 = cursor1.cursor()
    db_cursor9.execute(
        "SELECT COUNT(*) FROM forms_osd WHERE DATE(date_time) >= %s AND DATE(date_time) <= %s;", (form, to))
    result9 = db_cursor9.fetchone()
    db_cursor9.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor5 = cursor1.cursor()
    db_cursor5.execute(
        "SELECT COUNT(*) FROM reports WHERE DATE(date_time) >= %s AND DATE(date_time) <= %s AND course = %s", (form, to, "COE",))
    result5 = db_cursor5.fetchone()
    db_cursor5.close

    dcnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor6 = cursor1.cursor()
    db_cursor6.execute(
        "SELECT COUNT(*) FROM reports WHERE DATE(date_time) >= %s AND DATE(date_time) <= %s AND course = %s", (form, to, "CICS",))
    result6 = db_cursor6.fetchone()
    db_cursor6.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor7 = cursor1.cursor()
    db_cursor7.execute(
        "SELECT COUNT(*) FROM reports WHERE DATE(date_time) >= %s AND DATE(date_time) <= %s AND course = %s", (form, to, "CAFAD",))
    result7 = db_cursor7.fetchone()
    db_cursor7.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor8 = cursor1.cursor()
    db_cursor8.execute(
        "SELECT COUNT(*) FROM reports WHERE DATE(date_time) >= %s AND DATE(date_time) <= %s AND course = %s", (form, to, "CIT",))
    result8 = db_cursor8.fetchone()
    db_cursor8.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor10 = cursor1.cursor()
    db_cursor10.execute(
        "SELECT COUNT(*) FROM forms_osd WHERE DATE(date_time) >= %s AND DATE(date_time) <= %s AND course = %s", (form, to, "COE",))
    result10 = db_cursor10.fetchone()
    db_cursor10.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor11 = cursor1.cursor()
    db_cursor11.execute(
        "SELECT COUNT(*) FROM forms_osd WHERE DATE(date_time) >= %s AND DATE(date_time) <= %s AND course = %s", (form, to, "CICS",))
    result11 = db_cursor11.fetchone()
    db_cursor11.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor12 = cursor1.cursor()
    db_cursor12.execute(
        "SELECT COUNT(*) FROM forms_osd WHERE DATE(date_time) >= %s AND DATE(date_time) <= %s AND course = %s", (form, to, "CAFAD",))
    result12 = db_cursor12.fetchone()
    db_cursor12.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor13 = cursor1.cursor()
    db_cursor13.execute(
        "SELECT COUNT(*) FROM forms_osd WHERE DATE(date_time) >= %s AND DATE(date_time) <= %s AND course = %s", (form, to, "CIT",))
    result13 = db_cursor13.fetchone()
    db_cursor13.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor14 = cursor1.cursor()
    db_cursor14.execute(
        "SELECT COUNT(*) FROM reports WHERE DATE(date_time) >= %s AND DATE(date_time) <= %s AND course = %s AND status = %s", (form, to, "COE", "Rejected"))
    result14 = db_cursor14.fetchone()
    db_cursor14.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor15 = cursor1.cursor()
    db_cursor15.execute(
        "SELECT COUNT(*) FROM reports WHERE DATE(date_time) >= %s AND DATE(date_time) <= %s AND course = %s AND status = %s", (form, to, "COE", "Case Closed"))
    result15 = db_cursor15.fetchone()
    db_cursor15.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor16 = cursor1.cursor()
    db_cursor16.execute(
        "SELECT COUNT(*) FROM reports WHERE DATE(date_time) >= %s AND DATE(date_time) <= %s AND course = %s AND status = %s", (form, to, "CICS", "Rejected"))
    result16 = db_cursor16.fetchone()
    db_cursor16.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor17 = cursor1.cursor()
    db_cursor17.execute(
        "SELECT COUNT(*) FROM reports WHERE DATE(date_time) >= %s AND DATE(date_time) <= %s AND course = %s AND status = %s", (form, to, "CICS", "Case Closed"))
    result17 = db_cursor17.fetchone()
    db_cursor17.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor18 = cursor1.cursor()
    db_cursor18.execute(
        "SELECT COUNT(*) FROM reports WHERE DATE(date_time) >= %s AND DATE(date_time) <= %s AND course = %s AND status = %s ", (form, to, "CAFAD", "Rejected"))
    result18 = db_cursor18.fetchone()
    db_cursor18.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor19 = cursor1.cursor()
    db_cursor19.execute("SELECT COUNT(*) FROM reports WHERE DATE(date_time) >= %s AND DATE(date_time) <= %s AND course = %s AND status = %s",
                        (form, to, "CAFAD", "Case Closed"))
    result19 = db_cursor19.fetchone()
    db_cursor19.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor20 = cursor1.cursor()
    db_cursor20.execute(
        "SELECT COUNT(*) FROM reports WHERE DATE(date_time) >= %s AND DATE(date_time) <= %s AND course = %s AND status = %s", (form, to, "CIT", "Rejected"))
    result20 = db_cursor20.fetchone()
    db_cursor20.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor21 = cursor1.cursor()
    db_cursor21.execute(
        "SELECT COUNT(*) FROM reports WHERE DATE(date_time) >= %s AND DATE(date_time) <= %s AND course = %s AND status = %s", (form, to, "CIT", "Case Closed"))
    result21 = db_cursor21.fetchone()
    db_cursor21.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor22 = cursor1.cursor()
    db_cursor22.execute(
        "SELECT COUNT(*) FROM reports WHERE DATE(date_time) >= %s AND DATE(date_time) <= %s AND status = %s", (form, to, "Rejected",))
    result22 = db_cursor22.fetchone()
    db_cursor22.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor23 = cursor1.cursor()
    db_cursor23.execute(
        "SELECT COUNT(*) FROM reports WHERE DATE(date_time) >= %s AND DATE(date_time) <= %s AND status = %s", (form, to, "Case Closed",))
    result23 = db_cursor23.fetchone()
    db_cursor23.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor24 = cursor1.cursor()
    db_cursor24.execute(
        "SELECT COUNT(*) FROM forms_osd WHERE DATE(date_time) >= %s AND DATE(date_time) <= %s AND course = %s AND status = %s", (form, to, "COE", "Rejected"))
    result24 = db_cursor24.fetchone()
    db_cursor24.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor25 = cursor1.cursor()
    db_cursor25.execute(
        "SELECT COUNT(*) FROM forms_osd WHERE DATE(date_time) >= %s AND DATE(date_time) <= %s AND course = %s AND status = %s", (form, to, "COE", "Approved"))
    result25 = db_cursor25.fetchone()
    db_cursor25.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor26 = cursor1.cursor()
    db_cursor26.execute(
        "SELECT COUNT(*) FROM forms_osd WHERE DATE(date_time) >= %s AND DATE(date_time) <= %s AND course = %s AND status = %s", (form, to, "CICS", "Rejected"))
    result26 = db_cursor26.fetchone()
    db_cursor26.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor27 = cursor1.cursor()
    db_cursor27.execute(
        "SELECT COUNT(*) FROM forms_osd WHERE DATE(date_time) >= %s AND DATE(date_time) <= %s AND course = %s AND status = %s", (form, to, "CICS", "Approved"))
    result27 = db_cursor27.fetchone()
    db_cursor27.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor28 = cursor1.cursor()
    db_cursor28.execute(
        "SELECT COUNT(*) FROM forms_osd WHERE DATE(date_time) >= %s AND DATE(date_time) <= %s AND course = %s AND status = %s", (form, to, "CAFAD", "Rejected"))
    result28 = db_cursor28.fetchone()
    db_cursor28.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor29 = cursor1.cursor()
    db_cursor29.execute(
        "SELECT COUNT(*) FROM forms_osd WHERE DATE(date_time) >= %s AND DATE(date_time) <= %s AND course = %s AND status = %s", (form, to, "CAFAD", "Approved"))
    result29 = db_cursor29.fetchone()
    db_cursor29.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor30 = cursor1.cursor()
    db_cursor30.execute(
        "SELECT COUNT(*) FROM forms_osd WHERE DATE(date_time) >= %s AND DATE(date_time) <= %s AND course = %s AND status = %s", (form, to, "CIT", "Rejected"))
    result30 = db_cursor30.fetchone()
    db_cursor30.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor31 = cursor1.cursor()
    db_cursor31.execute(
        "SELECT COUNT(*) FROM forms_osd WHERE DATE(date_time) >= %s AND DATE(date_time) <= %s AND course = %s AND status = %s", (form, to, "CIT", "Approved"))
    result31 = db_cursor31.fetchone()
    db_cursor31.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor32 = cursor1.cursor()
    db_cursor32.execute(
        "SELECT COUNT(*) FROM forms_osd WHERE DATE(date_time) >= %s AND DATE(date_time) <= %s AND status = %s", (form, to, "Rejected"))
    result32 = db_cursor32.fetchone()
    db_cursor32.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor33 = cursor1.cursor()
    db_cursor33.execute(
        "SELECT COUNT(*) FROM forms_osd WHERE DATE(date_time) >= %s AND DATE(date_time) <= %s AND status = %s", (form, to, "Approved"))
    result33 = db_cursor33.fetchone()
    db_cursor33.close

    counts4={}
    counts5={}


    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor34 = cursor1.cursor()

    minor_offenses_sum = 0

    for status in minor:
  
        query = "SELECT COUNT(*) FROM sanctions WHERE sanction = %s"
        db_cursor34.execute(query, (status,))

        result4 = db_cursor34.fetchone()
        counts4[status] = str(result4[0])
        minor_offenses_sum += int(counts4[status])

    db_cursor34.close()

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor35 = cursor1.cursor()

    major_offenses_sum = 0

    for status in major:
  
        query = "SELECT COUNT(*) FROM sanctions WHERE sanction = %s"
        db_cursor35.execute(query, (status,))

        result5 = db_cursor35.fetchone()
        counts5[status] = str(result5[0])
        major_offenses_sum += int(counts5[status])

    db_cursor35.close()

    countreports = str(result[0])

    count1 = str(result5[0])
    count2 = str(result6[0])
    count3 = str(result7[0])
    count4 = str(result8[0])
    countrequest = str(result9[0])
    request1 = str(result10[0])
    request2 = str(result11[0])
    request3 = str(result12[0])
    request4 = str(result13[0])
    reject1 = str(result14[0])
    reject2 = str(result16[0])
    reject3 = str(result18[0])
    reject4 = str(result20[0])
    closed1 = str(result15[0])
    closed2 = str(result17[0])
    closed3 = str(result19[0])
    closed4 = str(result21[0])
    totalreject = str(result22[0])
    totalcaseclosed = str(result23[0])
    creject1 = str(result24[0])
    creject2 = str(result26[0])
    creject3 = str(result28[0])
    creject4 = str(result30[0])
    cclosed1 = str(result25[0])
    cclosed2 = str(result27[0])
    cclosed3 = str(result29[0])
    cclosed4 = str(result31[0])
    ctotalreject = str(result32[0])
    ctotalcaseclosed = str(result33[0])

    pdf_filename = 'reports.docx'
    doc = Document(pdf_filename)

    replace_placeholder(doc, "(date)", form)
    replace_placeholder(doc, "(date1)", to)
    replace_table_cell_placeholder1(doc.tables[0], 1, 1, count1, "coe")
    replace_table_cell_placeholder1(doc.tables[0], 2, 1, count2, "CICS")
    replace_table_cell_placeholder1(doc.tables[0], 3, 1, count3, "cafad")
    replace_table_cell_placeholder1(doc.tables[0], 4, 1, count4, "cit")
    replace_table_cell_placeholder1(doc.tables[0], 5, 1, countreports, "total")
    replace_table_cell_placeholder1(doc.tables[0], 1, 2, request1, "(coe1)")
    replace_table_cell_placeholder1(doc.tables[0], 2, 2, request2, "(cics1)")
    replace_table_cell_placeholder1(doc.tables[0], 3, 2, request3, "(cafad1)")
    replace_table_cell_placeholder1(doc.tables[0], 4, 2, request4, "(cit1)")
    replace_table_cell_placeholder1(
        doc.tables[0], 5, 2, countrequest, "(total1)")

    replace_table_cell_placeholder1(doc.tables[1], 1, 1, count1, "(coe2)")
    replace_table_cell_placeholder1(doc.tables[1], 2, 1, count2, "CICS")
    replace_table_cell_placeholder1(doc.tables[1], 3, 1, count3, "CAFAD")
    replace_table_cell_placeholder1(doc.tables[1], 4, 1, count4, "CIT")
    replace_table_cell_placeholder1(doc.tables[1], 5, 1, countreports, "(total)")

    replace_table_cell_placeholder1(doc.tables[1], 1, 2, reject1, "(coe2)")
    replace_table_cell_placeholder1(doc.tables[1], 2, 2, reject2, "CICS")
    replace_table_cell_placeholder1(doc.tables[1], 3, 2, reject3, "CAFAD")
    replace_table_cell_placeholder1(doc.tables[1], 4, 2, reject4, "CIT")
    replace_table_cell_placeholder1(doc.tables[1], 5, 2, totalreject, "(total)")

    replace_table_cell_placeholder1(doc.tables[1], 1, 3, closed1, "(coe2)")
    replace_table_cell_placeholder1(doc.tables[1], 2, 3, closed2, "CICS")
    replace_table_cell_placeholder1(doc.tables[1], 3, 3, closed3, "CAFAD")
    replace_table_cell_placeholder1(doc.tables[1], 4, 3, closed4, "CIT")
    replace_table_cell_placeholder1(
        doc.tables[1], 5, 3, totalcaseclosed, "(total)")

    replace_table_cell_placeholder1(doc.tables[2], 1, 1, request1, "(coe2)")
    replace_table_cell_placeholder1(doc.tables[2], 2, 1, request2, "CICS")
    replace_table_cell_placeholder1(doc.tables[2], 3, 1, request3, "CAFAD")
    replace_table_cell_placeholder1(doc.tables[2], 4, 1, request4, "CIT")
    replace_table_cell_placeholder1(
        doc.tables[2], 5, 1, countrequest, "(total)")
    replace_table_cell_placeholder1(doc.tables[2], 1, 2, creject1, "(coe2)")
    replace_table_cell_placeholder1(doc.tables[2], 2, 2, creject2, "CICS")
    replace_table_cell_placeholder1(doc.tables[2], 3, 2, creject3, "CAFAD")
    replace_table_cell_placeholder1(doc.tables[2], 4, 2, creject4, "CIT")
    replace_table_cell_placeholder1(
        doc.tables[2], 5, 2, ctotalreject, "(total)")
    replace_table_cell_placeholder1(doc.tables[2], 1, 3, cclosed1, "(coe2)")
    replace_table_cell_placeholder1(doc.tables[2], 2, 3, cclosed2, "CICS")
    replace_table_cell_placeholder1(doc.tables[2], 3, 3, cclosed3, "CAFAD")
    replace_table_cell_placeholder1(doc.tables[2], 4, 3, cclosed4, "CIT")
    replace_table_cell_placeholder1(doc.tables[2], 5, 3, ctotalcaseclosed, "(total)")



    replace_table_cell_placeholder1(doc.tables[3], 1, 1, counts4["12.1.1 - attendance, punctuality, cutting classes"], "coe")
    replace_table_cell_placeholder1(doc.tables[3], 2, 1, counts4["12.1.2 - dress code, uniform"], "CICS")
    replace_table_cell_placeholder1(doc.tables[3], 3, 1, counts4["12.1.3 - property misuse"], "cafad")
    replace_table_cell_placeholder1(doc.tables[3], 4, 1, counts4["12.1.4 - noise disturbance"], "cit")
    replace_table_cell_placeholder1(doc.tables[3], 5, 1, counts4["12.1.5 - posting violation"], "total")
    replace_table_cell_placeholder1(doc.tables[3], 6, 1, counts4["12.1.6 - notice removal"], "(coe1)")
    replace_table_cell_placeholder1(doc.tables[3], 7, 1, counts4["12.1.7 - littering"], "(cics1)")
    replace_table_cell_placeholder1(doc.tables[3], 8, 1, counts4["12.1.8 - smoking violation"], "(cafad1)")
    replace_table_cell_placeholder1(doc.tables[3], 9, 1, counts4["12.1.9 - trespassing"], "(cit1)")
    replace_table_cell_placeholder1(doc.tables[3], 10, 1, counts4["12.1.10 - misconduct"], "1")
    replace_table_cell_placeholder1(doc.tables[3], 11, 1, counts4["12.1.11 - harassment"], "2")
    replace_table_cell_placeholder1(doc.tables[3], 12, 1, counts4["12.1.12 - provocation, fight"], "3")
    replace_table_cell_placeholder1(doc.tables[3], 13, 1, counts4["12.1.13 - PDA"], "4")
    replace_table_cell_placeholder1(doc.tables[3], 14, 1, counts4["12.1.14 - truancy"], "5")
    replace_table_cell_placeholder1(doc.tables[3], 15, 1, str(minor_offenses_sum), "6")

    replace_table_cell_placeholder1(doc.tables[4], 1, 1, counts5["13.1 - repeat offenses"], "1")
    replace_table_cell_placeholder1(doc.tables[4], 2, 1, counts5["13.2 - insubordination"], "2")
    replace_table_cell_placeholder1(doc.tables[4], 3, 1, counts5["13.3 - smoking violation"], "3")
    replace_table_cell_placeholder1(doc.tables[4], 4, 1, counts5["13.4 - alcohol violation"], "4")
    replace_table_cell_placeholder1(doc.tables[4], 5, 1, counts5["13.5 - intoxication"], "5")
    replace_table_cell_placeholder1(doc.tables[4], 6, 1, counts5["13.6 - trespassing"], "6")
    replace_table_cell_placeholder1(doc.tables[4], 7, 1, counts5["13.7 - property misuse"], "7")
    replace_table_cell_placeholder1(doc.tables[4], 8, 1, counts5["13.8 - Reckless endangerment"], "8")
    replace_table_cell_placeholder1(doc.tables[4], 9, 1, counts5["13.9 - Gambling"], "9")
    replace_table_cell_placeholder1(doc.tables[4], 10, 1, counts5["13.10 - Identity fraud"], "10")
    replace_table_cell_placeholder1(doc.tables[4], 11, 1, counts5["13.11 - Misuse of university name/logo"], "11")
    replace_table_cell_placeholder1(doc.tables[4], 12, 1, counts5["13.12 - Unauthorized representation"], "12")
    replace_table_cell_placeholder1(doc.tables[4], 13, 1, counts5["13.13 - abusive behavior"], "13")
    replace_table_cell_placeholder1(doc.tables[4], 14, 1, counts5["13.14 - unauthorized membership"], "14")
    replace_table_cell_placeholder1(doc.tables[4], 15, 1, counts5["13.15 - online misconduct"], "15")
    replace_table_cell_placeholder1(doc.tables[4], 16, 1, counts5["13.16 - vandalism"], "16")
    replace_table_cell_placeholder1(doc.tables[4], 17, 1, counts5["13.17 - academic disruption"], "17")
    replace_table_cell_placeholder1(doc.tables[4], 18, 1, counts5["13.18 - solicitation"], "18")
    replace_table_cell_placeholder1(doc.tables[4], 19, 1, counts5["13.19 - physical harm"], "19")
    replace_table_cell_placeholder1(doc.tables[4], 20, 1, counts5["13.20 - weapons possession"], "20")
    replace_table_cell_placeholder1(doc.tables[4], 21, 1, counts5["13.21 - theft"], "21")
    replace_table_cell_placeholder1(doc.tables[4], 22, 1, counts5["13.22 - bribery"], "22")
    replace_table_cell_placeholder1(doc.tables[4], 23, 1, counts5["13.23 - sexual misconduct"], "23")
    replace_table_cell_placeholder1(doc.tables[4], 24, 1, counts5["13.24 - obscenity"], "24")
    replace_table_cell_placeholder1(doc.tables[4], 25, 1, counts5["13.25 - defamation"], "25")
    replace_table_cell_placeholder1(doc.tables[4], 26, 1, counts5["13.26 - physical harm"], "26")
    replace_table_cell_placeholder1(doc.tables[4], 27, 1, counts5["13.27 - falsification"], "27")
    replace_table_cell_placeholder1(doc.tables[4], 28, 1, counts5["13.28 - disrepute"], "28")
    replace_table_cell_placeholder1(doc.tables[4], 29, 1, counts5["13.29 - riot"], "29")
    replace_table_cell_placeholder1(doc.tables[4], 30, 1, counts5["13.30 - destruction of property"], "30")
    replace_table_cell_placeholder1(doc.tables[4], 31, 1, counts5["13.31 - burglary"], "31")
    replace_table_cell_placeholder1(doc.tables[4], 32, 1, counts5["13.32 - hazing"], "32")
    replace_table_cell_placeholder1(doc.tables[4], 33, 1, counts5["13.33 - drugs"], "33")
    replace_table_cell_placeholder1(doc.tables[4], 34, 1, counts5["13.34 - firearms possession"], "34")
    replace_table_cell_placeholder1(doc.tables[4], 35, 1, counts5["13.35 - threats"], "35")
    replace_table_cell_placeholder1(doc.tables[4], 36, 1, counts5["13.36 - felonies"], "36")
    replace_table_cell_placeholder1(doc.tables[4], 37, 1, counts5["13.37 - moral turpitude"], "37")
    replace_table_cell_placeholder1(doc.tables[4], 38, 1, counts5["14.1 - cheating, mobile phone"], "38")
    replace_table_cell_placeholder1(doc.tables[4], 39, 1, counts5["14.2 - cheating, talking"], "39")
    replace_table_cell_placeholder1(doc.tables[4], 40, 1, counts5["14.3 - cheating, dictating answers"], "40")
    replace_table_cell_placeholder1(doc.tables[4], 41, 1, counts5["14.4 - cheating, notes possession"], "41")
    replace_table_cell_placeholder1(doc.tables[4], 42, 1, counts5["14.5 - cheating, outside information"], "42")
    replace_table_cell_placeholder1(doc.tables[4], 43, 1, counts5["14.6 - cheating, leakage facilitation"], "43")
    replace_table_cell_placeholder1(doc.tables[4], 44, 1, counts5["14.7 - cheating, buying/selling questions"], "44")
    replace_table_cell_placeholder1(doc.tables[4], 45, 1, counts5["14.8 - cheating, copying answers"], "45")
    replace_table_cell_placeholder1(doc.tables[4], 46, 1, counts5["14.9 - cheating, covert devices"], "46")
    replace_table_cell_placeholder1(doc.tables[4], 47, 1, counts5["14.10 - cheating, impersonation"], "47")
    replace_table_cell_placeholder1(doc.tables[4], 48, 1, counts5["14.11 - plagiarism"], "48")
    replace_table_cell_placeholder1(doc.tables[4], 49, 1, counts5["14.12 - cheating, surrogate attendance"], "49")
    replace_table_cell_placeholder1(doc.tables[4], 50, 1, counts5["14.13 - plagiarism"], "50")
    replace_table_cell_placeholder1(doc.tables[4], 51, 1, counts5["14.14 - cheating, caught"], "51")
    replace_table_cell_placeholder1(doc.tables[4], 52, 1, counts5["14.15 - cheating, aiding"], "52")
    replace_table_cell_placeholder1(doc.tables[4], 53, 1, str(major_offenses_sum), "53")


    doc.save("modified_document.docx")
   

    pdfpath = os.path.join('modified_document.docx')

    file_name = f'{random_code}_Reports.docx'
    with open(pdfpath, "rb") as pdf_file:
        pdf_data = pdf_file.read()

    response = Response(pdf_data, content_type='application/octet-stream')
    response.headers['Content-Disposition'] = f'attachment; filename="{file_name}".docx'

    flash('The report is submitted', 'success')

    return response


@app.route('/submit_report', methods=['GET', 'POST'])
def submit_report():
    role = request.form.get('role')
    kind = request.form.get('forms')
    print(role)
    print(kind)
    print("test")
    if kind == "Formal Complaint":

        if role == "coord":
            course1 = session.get('courseall', '')

        else:
            course1 = session.get('course1', '')


        
        print("test1")
        department = request.form.get('department')
        provision = ""
        final = request.form.get('final')
        report_text = request.form.get('narrate')
        name = request.form.get('name')
        section = request.form.get('section')
        number = request.form.get('number')
        email = request.form.get('email')
        namecomplain = request.form.get('namecomplain')
        witness1 = request.form.get('witness1')
        witness2 = request.form.get('witness2')
        witness3 = request.form.get('witness3')
        evidence1 = request.form.get('evi1')
        evidence2 = request.form.get('evi2')
        evidence3 = request.form.get('evi3')
        pic = request.files['file7']
        current_datetime = datetime.now()
        current_date = current_datetime.date()
        formatted_date = current_date.strftime("%m/%d/%Y")
        random_code = generate_random_code()



        if pic is None:
            pic = ""


        if department == "CAFAD":
            Name_Coordinator = "Paula Joyce A. Buisan"

        elif department == "CICS":
            Name_Coordinator = "Lovely Rose Tipan Hernandez"


        elif department == "CIT":
            Name_Coordinator = "Dolfus G. Miciano"


        elif department == "COE":
            Name_Coordinator = "Dolfus G. Miciano"

        elif department == "COE1":
            Name_Coordinator = "Therezia O. Conti"

        elif department == "COE2":
            Name_Coordinator = "Belen E. Bagui"

        username = session.get('username', '')
        print(username)

        pdf_filename = 'Formal Complaint Letter.docx'

        doc = Document(pdf_filename)
        # Replace placeholders

        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor = cursor1.cursor()
        db_cursor.execute(
            "SELECT * FROM accounts_cics WHERE Name = %s", (name,))
        result_cics = db_cursor.fetchone()

        db_cursor.execute(
            "SELECT * FROM accounts_cafad WHERE Name = %s", (name,))
        result_cafad = db_cursor.fetchone()

        db_cursor.execute(
            "SELECT * FROM accounts_coe WHERE Name = %s", (name,))
        result_coe = db_cursor.fetchone()

        db_cursor.execute(
            "SELECT * FROM accounts_cit WHERE Name = %s", (name,))
        result_cit = db_cursor.fetchone()

        db_cursor.close()


        if result_cics:
            department1 = "CICS"
        
        elif result_cafad:
            department1 = "CAFAD"

        elif result_coe:
            department1 = "COE"

        elif result_cit:
            department1 = "CIT"

        replace_table_cell_placeholder1(doc.tables[0], 2, 2, formatted_date, "(date)")
        replace_table_cell_placeholder1(doc.tables[0], 4, 1, Name_Coordinator, "NAME")
        replace_table_cell_placeholder1(doc.tables[0], 11, 8, name, "(student)")
        replace_table_cell_placeholder1(doc.tables[0], 12, 8, department1, "(college)")
        replace_table_cell_placeholder1(doc.tables[0], 13, 8, section, "(section)")

        replace_table_cell_placeholder1(doc.tables[0], 23, 3, report_text, "(narration)")
        replace_table_cell_placeholder1(doc.tables[0], 30, 3, final, "(final)")
        replace_table_cell_placeholder_with_image(doc.tables[0], 37, 18, pic, "lol")
        replace_table_cell_placeholder1(doc.tables[0], 37, 18, namecomplain, "(NAME)")
        replace_table_cell_placeholder1(doc.tables[0], 38, 18, number, "(number)")
        replace_table_cell_placeholder1(doc.tables[0], 39, 18, email, "(email)")
        replace_table_cell_placeholder1(doc.tables[0], 40, 6, witness1, "(witness1)")
        replace_table_cell_placeholder1(doc.tables[0], 41, 6, witness2, "(witness2)")
        replace_table_cell_placeholder1(doc.tables[0], 42, 6, witness3, "(witness3)")
        replace_table_cell_placeholder1(doc.tables[0], 44, 9, evidence1, "(evidence1)")
        replace_table_cell_placeholder1(doc.tables[0], 45, 9, evidence2, "(evidence2)")
        replace_table_cell_placeholder1(doc.tables[0], 46, 9, evidence3, "(evidence3)")


        doc.save("modified_document.docx")

       
        pdfpath = os.path.join('modified_document.docx')

     

        
        file_name = f'{random_code}_Formal Complaint Letter'
        with open(pdfpath, "rb") as pdf_file:
            pdf_data = pdf_file.read()

        support_file = request.files['file4']
        support_file1 = request.files['file5']
        support_file2 = request.files['file6']

        # Check if the user submitted an empty supporting document file input

        if support_file.filename == '':
            support_data = None
            support_filename = "None"
            support_extension = "None"

        else:
            # Securely get the filenames and file extensions
            support_filename = secure_filename(support_file.filename)

            support_extension = os.path.splitext(support_filename)[1]

            # Read the file data into memory

            support_data = support_file.read()


        if support_file1.filename == '':
            support_data1 = None
            support_filename1 = "None"
            support_extension1 = "None"

        else:
            # Securely get the filenames and file extensions
            support_filename1 = secure_filename(support_file1.filename)

            support_extension1 = os.path.splitext(support_filename1)[1]

            # Read the file data into memory

            support_data1 = support_file1.read()

        if support_file2.filename == '':
            support_data2 = None
            support_filename2 = "None"
            support_extension2 = "None"

        else:
            # Securely get the filenames and file extensions
            support_filename2 = secure_filename(support_file2.filename)

            support_extension2 = os.path.splitext(support_filename2)[1]

            # Read the file data into memory

            support_data2 = support_file2.read()



        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor = cursor1.cursor()
        db_cursor.execute("INSERT INTO reports (report_id, course, report, file_form,file_form_name,file_support_name, file_support_type, file_support,file_support_name1, file_support_type1, file_support1, file_support_name2, file_support_type2, file_support2, username, date_time, status, course1) VALUES (%s,%s, %s,%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,%s)",
                            (random_code, department, report_text, pdf_data, file_name, support_filename, support_extension, support_data, support_filename1, support_extension1, support_data1,support_filename2, support_extension2, support_data2, username, current_datetime, "Pending", course1))
        cursor1.commit()

        db_cursor.close()

        flash('The report is submitted', 'success')

        if role == "coord":
            return redirect('/head')
        else:
            return redirect('/hello')

        
            
    else:
        if role == "coord":
            course1 = session.get('courseall', '')

        else:
            course1 = session.get('course1', '')

        department = request.form.get('department')
        remarks = request.form.get('remarks')
        report_text = request.form.get('Incident')
        name = request.form.get('name1')
        print("lol")
        section = request.form.get('section1')
        designation = request.form.get('designation')
        program = request.form.get('program')
        namecomplain = request.form.get('namecomplain')
        pic = request.files['file3']
        current_datetime = datetime.now()
        current_date = current_datetime.date()
        formatted_date = current_date.strftime("%m/%d/%Y")
        current_time = current_datetime.strftime('%I:%M %p')
        random_code = generate_random_code()

        username = session.get('username', '')

        pdf_filename = 'Incident Report.docx'
        doc = Document(pdf_filename)
        # Replace placeholders

        replace_table_cell_placeholder(doc.tables[0], 2, 3, str(current_date))
        replace_table_cell_placeholder(doc.tables[0], 3, 3, name)
        replace_table_cell_placeholder(doc.tables[0], 4, 3, department)
        replace_table_cell_placeholder(doc.tables[0], 5, 3, program)
        replace_table_cell_placeholder(doc.tables[0], 6, 4, report_text)
        replace_table_cell_placeholder(doc.tables[0], 10, 4, remarks)

        replace_table_cell_placeholder_with_image(
            doc.tables[0], 14, 1, pic, "(signature)", 2)
        replace_table_cell_placeholder1(
            doc.tables[0], 14, 1, namecomplain, "Amazing")
        replace_table_cell_placeholder1(
            doc.tables[0], 14, 1, designation, "(designation)")
        replace_table_cell_placeholder1(
            doc.tables[0], 14, 1, str(current_date), "lol")
        replace_table_cell_placeholder(doc.tables[0], 2, 8, current_time)
        replace_table_cell_placeholder(doc.tables[0], 3, 10, username)
        replace_table_cell_placeholder(doc.tables[0], 5, 8, section)

        replace_table_cell_placeholder(doc.tables[1], 2, 3, str(current_date))
        replace_table_cell_placeholder(doc.tables[1], 3, 3, name)
        replace_table_cell_placeholder(doc.tables[1], 4, 3, department)
        replace_table_cell_placeholder(doc.tables[1], 5, 3, program)
        replace_table_cell_placeholder(doc.tables[1], 6, 4, report_text)
        replace_table_cell_placeholder(doc.tables[1], 10, 4, remarks)

        replace_table_cell_placeholder_with_image(
            doc.tables[1], 14, 1, pic, "(signature)", 29)
        replace_table_cell_placeholder1(
            doc.tables[1], 14, 1, namecomplain, "Amazing")
        replace_table_cell_placeholder1(
            doc.tables[1], 14, 1, designation, "(designation)")
        replace_table_cell_placeholder1(
            doc.tables[1], 14, 1, str(current_date), "lol")
        replace_table_cell_placeholder(doc.tables[1], 2, 8, current_time)
        replace_table_cell_placeholder(doc.tables[1], 3, 10, username)
        replace_table_cell_placeholder(doc.tables[1], 5, 8, section)

        doc.save("modified_document.docx")

        pdfpath = os.path.join('modified_document.docx')

        
        file_name = f'{random_code}_Incident Report Letter'
        with open(pdfpath, "rb") as pdf_file:
            pdf_data = pdf_file.read()

        # Check if the POST request has the file part for the supporting document file
        if 'file4' not in request.files:
            flash('No supporting document file part')
            return redirect(request.url)

        support_file = request.files['file4']

        # Check if the user submitted an empty supporting document file input

        if support_file.filename == '':
            support_data = None
            support_filename = "None"
            support_extension = "None"

            cnx = create_connection_pool()
            cursor1=cnx.get_connection()
            db_cursor = cursor1.cursor()
            db_cursor.execute("INSERT INTO reports (report_id, course, report, file_form,file_form_name,file_support_name, file_support_type, file_support, username, date_time, status,course1) VALUES (%s,%s, %s,%s, %s, %s, %s, %s, %s, %s, %s,%s)",
                              (random_code, department, report_text, pdf_data, file_name, support_filename, support_extension, support_data, username, current_datetime, "Pending",course1))
            cursor1.commit()

            db_cursor.close()

            flash('The report is submitted', 'success')
            if role == "coord":
                return redirect('/head')
            else:
                return redirect('/hello')
        else:

            support_filename = secure_filename(support_file.filename)

            support_extension = os.path.splitext(support_filename)[1]

            # Read the file data into memory

            support_data = support_file.read()

            # Insert the report with file information into the database, including file data
            cnx = create_connection_pool()
            cursor1=cnx.get_connection()
            db_cursor = cursor1.cursor()
            db_cursor.execute("INSERT INTO reports (report_id, course, report, file_form, file_form_name, file_support_name, file_support_type, file_support, username, date_time, status,course1) VALUES (%s,%s, %s, %s, %s, %s, %s, %s, %s,%s,%s,%s)",
                              (random_code, department, report_text, pdf_data, file_name, support_filename, support_extension, support_data, username, current_datetime, "Pending",course1))
            cursor1.commit()

            db_cursor.close()

            flash('The report is submitted', 'success')
            if role == "coord":
                return redirect('/head')
            else:
                return redirect('/hello')


@app.route('/submit_request', methods=['GET', 'POST'])
def submit_request():
    kind = request.form.get('forms')
    print(kind)
    if kind == "Temporary Gate Pass":
       
        course1 = session.get('course1', '')

        remarks = request.form.get('remarks')
        print(remarks)
        department = request.form.get('department1')
        print(department)
        section = request.form.get('section2')
        program = request.form.get('program')
        current_datetime = datetime.now()
        random_code = generate_random_code()
        current_date = current_datetime.date()
        formatted_date = current_date.strftime("%m/%d/%Y")

        student = session.get('namestudent', '')
        username = session.get('username', '')


        pdf_filename = 'Temporary Gate Pass.docx'
        doc = Document(pdf_filename)

        replace_table_cell_placeholder1(doc.tables[0], 2, 11, formatted_date, "(date)")
        
        
        replace_table_cell_placeholder1(doc.tables[0], 3, 3, student, "(name)")
        replace_table_cell_placeholder1(doc.tables[0], 6, 4, remarks, "(remarks)")
        replace_table_cell_placeholder1(doc.tables[0], 3, 11, username, "(code)")
        replace_table_cell_placeholder1(doc.tables[0], 5, 9, section, "(section)")
        replace_table_cell_placeholder1(doc.tables[0], 4, 3, department, "(department)")
        replace_table_cell_placeholder1(doc.tables[0], 5, 3, program, "(program)")

        replace_table_cell_placeholder1( doc.tables[1], 2, 11, formatted_date, "(date)")
        
        replace_table_cell_placeholder1(doc.tables[1], 3, 3, student, "(name)")
        replace_table_cell_placeholder1(doc.tables[1], 6, 4, remarks, "(remarks)")
        replace_table_cell_placeholder1(doc.tables[1], 3, 11, username, "(code)")
        replace_table_cell_placeholder1(doc.tables[1], 5, 9, section, "(section)")
        replace_table_cell_placeholder1(doc.tables[1], 4, 3, department, "(department)")
        replace_table_cell_placeholder1(doc.tables[1], 5, 3, program, "(program)")

        doc.save("modified_document.docx")
       
        pdfpath = os.path.join('modified_document.docx')

        file_name = f'{random_code}_Temporary Gate Pass Letter'
        with open(pdfpath, "rb") as pdf_file:
            pdf_data = pdf_file.read()

        # Check if the POST request has the file part for the supporting document file
        if 'file5' not in request.files:
            flash('No supporting document file part')
            return redirect(request.url)

        support_file = request.files['file5']

        # Check if the user submitted an empty supporting document file input

        if support_file.filename == '':
            support_data = None
            support_filename = "None"
            support_extension = "None"

            cnx = create_connection_pool()
            cursor1=cnx.get_connection()
            db_cursor = cursor1.cursor()
            db_cursor.execute("INSERT INTO forms_osd (form_id,course,report,file_form_name,file_form_type, file_form, file_support_name, file_support_type, file_support, username, date_time, status,remarks,course1) VALUES (%s, %s,%s, %s, %s, %s, %s, %s, %s, %s, %s,%s,%s,%s)",
                              (random_code, department, remarks, file_name, pdf_data, ".pdf",support_filename, support_extension, support_data, username, current_datetime, "Pending",remarks,course1))
            cursor1.commit()

            db_cursor.close()

            flash('The report is submitted', 'success')
            return redirect('/hello')
        else:

            support_filename = secure_filename(support_file.filename)

            support_extension = os.path.splitext(support_filename)[1]

            # Read the file data into memory

            support_data = support_file.read()

            # Insert the report with file information into the database, including file data
            cnx = create_connection_pool()
            cursor1=cnx.get_connection()
            db_cursor = cursor1.cursor()
            db_cursor.execute("INSERT INTO forms_osd (form_id,course,report,file_form_name, file_form,file_form_type, file_support_name, file_support_type, file_support, username, date_time, status,remarks,course1) VALUES (%s, %s,%s, %s, %s, %s, %s, %s, %s, %s, %s,%s,%s,%s)",
                              (random_code, department, remarks, file_name, pdf_data, ".pdf",support_filename, support_extension, support_data, username, current_datetime, "Pending",remarks,course1))
            cursor1.commit()

            db_cursor.close()

            flash('The report is submitted', 'success')
            return redirect('/hello')

    elif kind == "Request for Non-Wearing of Uniform":
        
        course1 = session.get('course1', '')
        fieldwork = request.form.get('fieldwork')
        prolonged = request.form.get('prolonged')
        foreign = request.form.get('foreign')
        pregnant = request.form.get('pregnant')
        cases = request.form.get('cases')
        majeure = request.form.get('majeure')
        internship = request.form.get('internship')
        specify = request.form.get('specify')
        remarks = "The details of the report is located in the document"
        department = request.form.get('department1')
        specify1 = request.form.get('specifyTextarea')
        print(specify1)
        print(department)
        section = request.form.get('section2')
        college = request.form.get('college')
        program = request.form.get('program')
        current_datetime = datetime.now()
        random_code = generate_random_code()
        current_date = current_datetime.date()
        formatted_date = current_date.strftime("%m/%d/%Y")
        pic = request.files['file3']

        student = session.get('namestudent', '')
        username = session.get('username', '')

        if fieldwork == "fieldwork":
            status = "checked"
        else:
            status = "not"

        if prolonged == "prolonged":
            status1 = "checked"
        else:
            status1 = "not"

        if foreign == "foreign":
            status2 = "checked"
        else:
            status2 = "not"

        if pregnant == "pregnant":
            status3 = "checked"
        else:
            status3 = "not"

        if cases == "cases":
            status4 = "checked"
        else:
            status4 = "not"

        if majeure == "majeure":
            status5 = "checked"
        else:
            status5 = "not"

        if internship == "internship":
            status6 = "checked"
        else:
            status6 = "not"

        if specify == "specify":
            status7 = "checked"
        else:
            status7 = "not"

        pdf_filename = 'Request for Non-Wearing of Uniform.docx'
        doc = Document(pdf_filename)

        toggle_table_cell_checkbox(doc.tables[0], 6, 0, status)
        toggle_table_cell_checkbox(doc.tables[0], 7, 0, status1)
        toggle_table_cell_checkbox(doc.tables[0], 8, 0, status2)
        toggle_table_cell_checkbox(doc.tables[0], 9, 0, status3)
        toggle_table_cell_checkbox(doc.tables[0], 10, 0, status4)
        toggle_table_cell_checkbox(doc.tables[0], 11, 0, status5)
        toggle_table_cell_checkbox(doc.tables[0], 12, 0, status6)
        toggle_table_cell_checkbox(doc.tables[0], 13, 0, status7)

        replace_table_cell_placeholder1(doc.tables[0], 16, 1, formatted_date, "(date)")
        replace_table_cell_placeholder1(doc.tables[0], 13, 2, specify1, "(specify)")
        replace_table_cell_placeholder1(doc.tables[0], 2, 4, student, "(name)")
        replace_table_cell_placeholder1(doc.tables[0], 3, 4, department, "(college)")
        replace_table_cell_placeholder1(doc.tables[0], 4, 4, program, "(program)")
        replace_table_cell_placeholder1(doc.tables[0], 2, 10, username, "(srcode)")
        replace_table_cell_placeholder1(doc.tables[0], 4, 10, section, "(section)")
        replace_table_cell_placeholder_with_image(doc.tables[0], 16, 1, pic, "(signature)", 29)
        replace_table_cell_placeholder1(doc.tables[0], 16, 1, student, "(name)")

        doc.save("modified_document.docx")
        pdfpath = os.path.join('modified_document.docx')

        file_name = f'{random_code}_Request for Non-Wearing of Uniform'
        with open(pdfpath, "rb") as pdf_file:
            pdf_data = pdf_file.read()

        # Check if the POST request has the file part for the supporting document file
        if 'file6' not in request.files:
            flash('No supporting document file part')
            return redirect(request.url)

        support_file = request.files['file6']

        # Check if the user submitted an empty supporting document file input

        if support_file.filename == '':
            support_data = None
            support_filename = "None"
            support_extension = "None"

            cnx = create_connection_pool()
            cursor1=cnx.get_connection()
            db_cursor = cursor1.cursor()
            db_cursor.execute("INSERT INTO forms_osd (form_id,course,report,file_form_name, file_form,file_form_type, file_support_name, file_support_type, file_support, username, date_time, status,remarks,course1) VALUES (%s, %s,%s, %s, %s, %s, %s, %s, %s, %s, %s,%s,%s,%s)",
                              (random_code, department, remarks, file_name, pdf_data,".pdf", support_filename, support_extension, support_data, username, current_datetime, "Pending","",course1))
            cursor1.commit()

            db_cursor.close()

            flash('The report is submitted', 'success')
            return redirect('/hello')
        else:

            support_filename = secure_filename(support_file.filename)

            support_extension = os.path.splitext(support_filename)[1]

            # Read the file data into memory

            support_data = support_file.read()

            # Insert the report with file information into the database, including file data
            cnx = create_connection_pool()
            cursor1=cnx.get_connection()
            db_cursor = cursor1.cursor()
            db_cursor.execute("INSERT INTO forms_osd (form_id,course,report,file_form_name, file_form,file_form_type, file_support_name, file_support_type, file_support, username, date_time, status,remarks,course1) VALUES (%s, %s,%s, %s, %s, %s, %s, %s, %s, %s, %s,%s,%s,%s)",
                              (random_code, department, remarks, file_name, pdf_data,".pdf", support_filename, support_extension, support_data, username, current_datetime, "Pending","",course1))
            cursor1.commit()

            db_cursor.close()

            flash('The report is submitted', 'success')
            return redirect('/hello')

            # Request for new id
    elif kind == "Request for New ID":
     
        course1 = session.get('course1', '')
        fieldwork = request.form.get('fieldwork')
        prolonged = request.form.get('prolonged')
        foreign = request.form.get('foreign')
        pregnant = request.form.get('pregnant')
        specify2 = request.form.get('specify1')
        print(specify2)
        remarks = "The details of the report is located in the document"
        department = request.form.get('department1')
        print(department)
        specify3 = request.form.get('specifyTextarea1')
        section = request.form.get('section1')
        college = request.form.get('college')
        program = request.form.get('program')
        current_datetime = datetime.now()
        random_code = generate_random_code()
        current_date = current_datetime.date()
        formatted_date = current_date.strftime("%m/%d/%Y")
        pic = request.files['file8']
        print(pic)

        student = session.get('namestudent', '')
        print(student)
        username = session.get('username', '')

        if fieldwork == "fieldwork":
            status = "checked"
        else:
            status = "not"

        if prolonged == "prolonged":
            status1 = "checked"
        else:
            status1 = "not"

        if foreign == "foreign":
            status2 = "checked"
        else:
            status2 = "not"

        if pregnant == "pregnant":
            status3 = "checked"

        else:
            status3 = "not"

        if specify2 == "specify1":
            status7 = "checked"
        else:
            status7 = "not"


        pdf_filename = 'request for new id.docx'
        doc = Document(pdf_filename)
# problem
        replace_table_cell_placeholder2(doc.tables[0], 7, 0, status, "SHIFT")
        replace_table_cell_placeholder2(doc.tables[0], 7, 4, status1, "LOST")
        replace_table_cell_placeholder2(doc.tables[0], 7, 8, status2, "TORN")
        replace_table_cell_placeholder2(doc.tables[0], 9, 2, status3, "UPDATE")
        replace_table_cell_placeholder2(doc.tables[0], 9, 4, status7, "OTHERS")

        replace_table_cell_placeholder1(doc.tables[0], 2, 3, formatted_date, "(date)")
        replace_table_cell_placeholder1(doc.tables[0], 10, 1, formatted_date, "(date1)")
       
        replace_table_cell_placeholder1(doc.tables[0], 8, 8, specify3, "(specify)")
        replace_table_cell_placeholder1(doc.tables[0], 3, 3, student, "(name)")
        replace_table_cell_placeholder1(doc.tables[0], 4, 3, department, "(college)")
        replace_table_cell_placeholder1(doc.tables[0], 5, 3, program, "(program)")
        replace_table_cell_placeholder1(doc.tables[0], 3, 10, username, "(srcode)")
        replace_table_cell_placeholder1(doc.tables[0], 5, 10, section, "(yearlevel)")
        replace_table_cell_placeholder_with_image(doc.tables[0], 10, 1, pic, "(signature)", 29)
        replace_table_cell_placeholder1(doc.tables[0], 10, 1, student, "(name)")

        replace_table_cell_placeholder2(doc.tables[1], 7, 0, status, "SHIFT")
        replace_table_cell_placeholder2(doc.tables[1], 7, 4, status1, "LOST")
        replace_table_cell_placeholder2(doc.tables[1], 7, 8, status2, "TORN")
        replace_table_cell_placeholder2(doc.tables[1], 9, 2, status3, "UPDATE")
        replace_table_cell_placeholder2(doc.tables[1], 9, 4, status7, "OTHERS")

        replace_table_cell_placeholder1(doc.tables[1], 2, 3, formatted_date, "(date)")
        replace_table_cell_placeholder1(doc.tables[1], 10, 1, formatted_date, "(date1)")
      
        replace_table_cell_placeholder1(doc.tables[1], 8, 8, specify3, "(specify)")
        replace_table_cell_placeholder1(doc.tables[1], 3, 3, student, "(name)")
        replace_table_cell_placeholder1(doc.tables[1], 4, 3, department, "(college)")
        replace_table_cell_placeholder1(doc.tables[1], 5, 3, program, "(program)")
        replace_table_cell_placeholder1(doc.tables[1], 3, 10, username, "(srcode)")
        replace_table_cell_placeholder1(doc.tables[1], 5, 10, section, "(yearlevel)")
        replace_table_cell_placeholder_with_image(doc.tables[1], 10, 1, pic, "(signature)", 29)
        replace_table_cell_placeholder1(doc.tables[1], 10, 1, student, "(name)")

        doc.save("modified_document.docx")
        pdfpath = os.path.join('modified_document.docx')

        file_name = f'{random_code}_Request for New ID'
        with open(pdfpath, "rb") as pdf_file:
            pdf_data = pdf_file.read()

        # Check if the POST request has the file part for the supporting document file
        if 'file4' not in request.files:
            flash('No supporting document file part')
            return redirect(request.url)

        support_file = request.files['file4']

        # Check if the user submitted an empty supporting document file input
        if support_file.filename == '':
            support_data = None
            support_filename = "None"
            support_extension = "None"

            cnx = create_connection_pool()
            cursor1=cnx.get_connection()
            db_cursor = cursor1.cursor()
            db_cursor.execute("INSERT INTO forms_osd (form_id,course,report,file_form_name, file_form,file_form_type, file_support_name, file_support_type, file_support, username, date_time, status,remarks,course1) VALUES (%s, %s,%s, %s, %s, %s, %s, %s, %s, %s, %s,%s,%s,%s)",
                              (random_code, department, remarks, file_name, pdf_data,".pdf", support_filename, support_extension, support_data, username, current_datetime, "Pending","",course1))
            cursor1.commit()

            db_cursor.close()

            flash('The report is submitted', 'success')
            return redirect('/hello')
        else:

            support_filename = secure_filename(support_file.filename)

            support_extension = os.path.splitext(support_filename)[1]

            # Read the file data into memory

            support_data = support_file.read()

            # Insert the report with file information into the database, including file data
            cnx = create_connection_pool()
            cursor1=cnx.get_connection()
            db_cursor = cursor1.cursor()
            db_cursor.execute("INSERT INTO forms_osd (form_id,course,report,file_form_name, file_form,file_form_type, file_support_name, file_support_type, file_support, username, date_time, status,remarks,course1) VALUES (%s, %s,%s, %s, %s, %s, %s, %s, %s, %s, %s,%s,%s,%s)",
                              (random_code, department, remarks, file_name, pdf_data,".pdf", support_filename, support_extension, support_data, username, current_datetime, "Pending","",course1))
            cursor1.commit()

            db_cursor.close()

            flash('The report is submitted', 'success')
            return redirect('/hello')


@app.route('/submit_call', methods=['POST'])
def submit_call():

    student = request.form.get('student')
    section = request.form.get('section')
    pic = request.files['file7']
    Time = request.form.get('meeting-time')
    # Parse the input time
    parsed_time = datetime.strptime(Time, "%H:%M")

    # Convert it to the desired format
    formatted_time = parsed_time.strftime("%I:%M %p")

    date2 = request.form.get('date2')
    date_format = '%Y-%m-%d'
    formatted_date1 = datetime.strptime(date2, date_format) 
    formatted_date_string = formatted_date1.strftime('%m/%d/%Y')
    
    remarks = request.form.get('remarks')
    current_datetime = datetime.now()
    random_code = generate_random_code()
    current_date = current_datetime.date()
    formatted_date = current_date.strftime("%m/%d/%Y")

    username = session.get('namestudent', '')

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()
    db_cursor.execute(
        "SELECT * FROM accounts_cics WHERE Username = %s", (student,))
    result_cics = db_cursor.fetchone()

    db_cursor.execute(
        "SELECT * FROM accounts_cafad WHERE Username = %s", (student,))
    result_cafad = db_cursor.fetchone()

    db_cursor.execute("SELECT * FROM accounts_coe WHERE Username = %s", (student,))
    result_coe = db_cursor.fetchone()

    db_cursor.execute("SELECT * FROM accounts_cit WHERE Username = %s", (student,))
    result_cit = db_cursor.fetchone()

    db_cursor.close()

    if result_cics:
        college = 'CICS'
        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor2 = cursor1.cursor()
        db_cursor2.execute(
            "SELECT Course FROM accounts_cics WHERE Username = %s", (student,))
        course1 = db_cursor2.fetchone()
        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor3 = cursor1.cursor()
        db_cursor3.execute(
            "SELECT Name FROM accounts_cics WHERE Username = %s", (student,))
        srcode1 = db_cursor3.fetchone()

        if course1:
            srcode = srcode1[0]
            # Get the first (and only) element of the tuple
            course = course1[0]
            print(course)  # Now, 'course' is a string
        else:
            print("No course found for the student.")

    elif result_cafad:
        college = 'CAFAD'
        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor2 = cursor1.cursor()
        db_cursor2.execute(
            "SELECT Course FROM accounts_cafad WHERE Username = %s", (student,))
        course1 = db_cursor2.fetchone()
        db_cursor2.close()
        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor3 = cursor1.cursor()
        db_cursor3.execute(
            "SELECT Name FROM accounts_cafad WHERE Username = %s", (student,))
        srcode1 = db_cursor3.fetchone()

        if course1:
            srcode = srcode1[0]
            # Get the first (and only) element of the tuple
            course = course1[0]
            print(course)  # Now, 'course' is a string
        else:
            print("No course found for the student.")

    elif result_coe:
        college = 'COE'
        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor2 = cursor1.cursor()
        db_cursor2.execute(
            "SELECT Course FROM accounts_coe WHERE Username = %s", (student,))
        course1 = db_cursor2.fetchone()
        db_cursor2.close()
        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor3 = cursor1.cursor()
        db_cursor3.execute(
            "SELECT Name FROM accounts_coe WHERE Username = %s", (student,))
        srcode1 = db_cursor3.fetchone()

        if course1:
            srcode = srcode1[0]
            # Get the first (and only) element of the tuple
            course = course1[0]
            print(course)  # Now, 'course' is a string
        else:
            print("No course found for the student.")

    elif result_cit:
        college = 'CIT'
        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor2 = cursor1.cursor()
        db_cursor2.execute(
            "SELECT Course FROM accounts_cit WHERE Username = %s", (student,))
        course1 = db_cursor2.fetchone()

        db_cursor2.close()

        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor3 = cursor1.cursor()
        db_cursor3.execute(
            "SELECT Name FROM accounts_cit WHERE Username = %s", (student,))
        srcode1 = db_cursor3.fetchone()

        if course1:
            srcode = srcode1[0]
            # Get the first (and only) element of the tuple
            course = course1[0]
            print(course)  # Now, 'course' is a string
        else:
            print("No course found for the student.")

    else:
        user_source = 'CAFAD'  # Handle the case where the user source is not found

    

    pdf_filename = 'call slip.docx'
    doc = Document(pdf_filename)

    replace_table_cell_placeholder1(doc.tables[0], 2, 3, srcode, "(name)")
    replace_table_cell_placeholder1(doc.tables[0], 4, 9, section, "(section)")
    replace_table_cell_placeholder1(
        doc.tables[0], 6, 6, formatted_time, "(time)")
    replace_table_cell_placeholder1(doc.tables[0], 6, 3, formatted_date_string, "(date1)")
    replace_table_cell_placeholder1(doc.tables[0], 3, 3, college, "(college)")
    replace_table_cell_placeholder1(doc.tables[0], 4, 3, course, "(program)")
    replace_table_cell_placeholder1(
        doc.tables[0], 2, 8, formatted_date, "(date)")
    replace_table_cell_placeholder1(doc.tables[0], 7, 1, username, "NAME")
    replace_table_cell_placeholder_with_image(doc.tables[0], 7, 1, pic, "(signature)")
    replace_table_cell_placeholder1(doc.tables[0], 7, 1, formatted_date, "DATE2")

    replace_table_cell_placeholder1(doc.tables[1], 2, 3, srcode, "(name)")
    replace_table_cell_placeholder1(doc.tables[1], 4, 9, section, "(section)")
    replace_table_cell_placeholder1(
        doc.tables[1], 6, 6, formatted_time, "(time)")
    replace_table_cell_placeholder1(doc.tables[1], 6, 3, formatted_date_string, "(date1)")
    replace_table_cell_placeholder1(doc.tables[1], 3, 3, college, "(college)")
    replace_table_cell_placeholder1(doc.tables[1], 4, 3, course, "(program)")
    replace_table_cell_placeholder1(
        doc.tables[1], 2, 8, formatted_date, "(date)")
    replace_table_cell_placeholder1(doc.tables[1], 7, 1, username, "NAME")
    replace_table_cell_placeholder_with_image(doc.tables[1], 7, 1, pic, "(signature)")
    replace_table_cell_placeholder1(doc.tables[1], 7, 1, formatted_date, "DATE2")

    replace_table_cell_placeholder1(doc.tables[2], 2, 3, srcode, "(name)")
    replace_table_cell_placeholder1(doc.tables[2], 4, 9, section, "(section)")
    replace_table_cell_placeholder1(
        doc.tables[2], 6, 6, formatted_time, "(time)")
    replace_table_cell_placeholder1(doc.tables[2], 6, 3, formatted_date_string, "(date1)")
    replace_table_cell_placeholder1(doc.tables[2], 3, 3, college, "(college)")
    replace_table_cell_placeholder1(doc.tables[2], 4, 3, course, "(program)")
    replace_table_cell_placeholder1(
        doc.tables[2], 2, 8, formatted_date, "(date)")
    replace_table_cell_placeholder1(doc.tables[2], 7, 1, username, "NAME")
    replace_table_cell_placeholder_with_image(doc.tables[2], 7, 1, pic, "(signature)")
    replace_table_cell_placeholder1(doc.tables[2], 7, 1, formatted_date, "DATE2")

    doc.save("modified_document.docx")
    pdfpath = os.path.join('modified_document.docx')

    file_name = f'{random_code}_Call Slip'
    with open(pdfpath, "rb") as pdf_file:
        pdf_data = pdf_file.read()

    notifs(srcode, "You have a new call slip")
    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor1 = cursor1.cursor()
    db_cursor1.execute("INSERT INTO callslip (call_id, name, coord,reason, date, time,file, file_name, date_issued) VALUES (%s, %s, %s, %s, %s,%s,%s,%s,%s)",
                       (random_code, srcode, username, remarks, formatted_date_string, formatted_time, pdf_data, file_name,formatted_date))
    cursor1.commit()
    
    db_cursor1.close()
    

    return redirect('/head')


@app.route('/submit_written', methods=['POST'])
def submit_written():

    kind = request.form.get('forms')
    print(kind)
    if kind == "Written Warning":

        remarks = request.form.get('remarks')
        complainant = request.form.get('student2')
        norms = request.form.get('norms')
        courseorposition = session.get('course', '')
        department = request.form.get('department')
        sanction = request.form.get('sanctions')
        students = request.form.get('student')

        date2 = request.form.get('date2')
        current_datetime = datetime.now()
        random_code = generate_random_code()
        current_date = current_datetime.date()
        formatted_date = current_date.strftime("/%m/%d/%Y")

        username = session.get('namestudent', '')

        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor1 = cursor1.cursor()
        db_cursor1.execute(
            "SELECT Name FROM accounts_cics WHERE Username = %s;", (students,))
        result_cics = db_cursor1.fetchone()

        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor1 = cursor1.cursor()
        db_cursor1.execute(
            "SELECT Name FROM accounts_cafad WHERE Username = %s;", (students,))
        result_cafad = db_cursor1.fetchone()

        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor1 = cursor1.cursor()
        db_cursor1.execute(
            "SELECT Name FROM accounts_coe WHERE Username = %s;", (students,))
        result_coe = db_cursor1.fetchone()

        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor1 = cursor1.cursor()
        db_cursor1.execute(
            "SELECT Name FROM accounts_cit WHERE Username = %s;", (students,))
        result_cit = db_cursor1.fetchone()

        if result_cics:
            srcode = students
            name = result_cics[0]

        elif result_cafad:
            srcode = students
            name = result_cafad[0]

        elif result_coe:
            srcode = result_coe[0]
            name = result_coe[0]

        elif result_cit:
            srcode = students
            name = result_cit[0]

        print(srcode+"lol")

        notifs(srcode, "You have a new sanction")

        db_cursor1.close()

        sanction_mapping = [
            "12.1.1 - attendance, punctuality, cutting classes",
            "12.1.2 - dress code, uniform",
            "12.1.3 - property misuse",
            "12.1.4 - noise disturbance",
            "12.1.5 - posting violation",
            "12.1.6 - notice removal",
            "12.1.7 - littering",
            "12.1.8 - smoking violation",
            "12.1.9 - trespassing",
            "12.1.10 - misconduct",
            "12.1.11 - harassment",
            "12.1.12 - provocation, fight",
            "12.1.13 - PDA",
            "12.1.14 - truancy",
        ]
        sanction_mapping1 = [
            "13.1 - repeat offenses",
            "13.2 - insubordination",
            "13.3 - smoking violation",
            "13.4 - alcohol violation",
            "13.5 - intoxication",
            "13.6 - trespassing",
            "13.7 - property misuse",
            "13.13 - abusive behavior",
            "13.14 - unauthorized membership",
            "13.15 - online misconduct",
            "13.16 - vandalism",
            "13.17 - academic disruption",
            "13.18 - solicitation",
            "13.19 - physical harm",
            "13.20 - weapons possession",
            "13.21 - theft",
            "13.22 - bribery",
            "13.23 - sexual misconduct",
            "13.24 - obscenity",
            "13.25 - defamation",
            "13.26 - physical harm",
            "13.27 - falsification",
            "13.28 - disrepute",
            "13.29 - riot",
            "13.30 - destruction of property",
            "13.31 - burglary",
            "13.32 - hazing",
            "13.33 - drugs",
            "13.34 - firearms possession",
            "13.35 - threats",
            "13.36 - felonies",
            "13.37 - moral turpitude",
        ]
        sanction_mapping2 = [
            "14.1 - cheating, mobile phone",
            "14.2 - cheating, talking",
            "14.3 - cheating, dictating answers",
            "14.4 - cheating, notes possession",
            "14.5 - cheating, outside information",
            "14.6 - cheating, leakage facilitation",
            "14.7 - cheating, buying/selling questions",
            "14.8 - cheating, copying answers",
            "14.9 - cheating, covert devices",
            "14.10 - cheating, impersonation",
            "14.11 - plagiarism",
            "14.12 - cheating, surrogate attendance",
            "14.13 - plagiarism",
            "14.14 - cheating, caught",
            "14.15 - cheating, aiding"
        ]

        sanction_number = None

        if sanction in sanction_mapping:
            sanction_number = "12"
        elif sanction in sanction_mapping1:
            sanction_number = "13"
        elif sanction in sanction_mapping2:
            sanction_number = "14"



        pdf_filename = 'written warning.docx'
        doc = Document(pdf_filename)

        replace_table_cell_placeholder1(doc.tables[0], 2, 12, formatted_date,"(date)")
        replace_table_cell_placeholder1(doc.tables[0], 3, 3, name,"(name)")
        replace_table_cell_placeholder1(doc.tables[0], 6, 7, date2,"(date2)")
        replace_table_cell_placeholder1(doc.tables[0], 7, 10, remarks,"(complain)")
        replace_table_cell_placeholder1(doc.tables[0], 6, 10, complainant,"(name1)")
        replace_table_cell_placeholder1(doc.tables[0], 11, 6, sanction_number,"(section)")
        replace_table_cell_placeholder1(doc.tables[0], 19, 8, username, "coord")
        replace_table_cell_placeholder1(doc.tables[0], 22, 2, username, "NAME")
        replace_table_cell_placeholder1(doc.tables[0], 12, 2, norms, "norms")

        doc.save("modified_document.docx")

        pdfpath = os.path.join('modified_document.docx')

        file_name = f'{random_code}_Written Warning'
        with open(pdfpath, "rb") as pdf_file:
            pdf_data = pdf_file.read()

        # Insert the report with file information into the database, including file data
        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor = cursor1.cursor()
        db_cursor.execute("INSERT INTO sanctions (sanctions_id,username, course, date_time, sanction, written, written_name,type) VALUES (%s,%s, %s, %s, %s, %s,%s,%s)",
                          (random_code, name, courseorposition, current_datetime, sanction, pdf_data, file_name, kind))
        cursor1.commit()
        db_cursor.close()




        flash('Report submitted successfully'
              )
        return redirect('/head')

    elif kind == 'Written Reprimand':

        remarks = request.form.get('remarks')
        norms = request.form.get('norms')
        courseorposition = session.get('course', '')
        print(courseorposition)
        department = request.form.get('department')
        sanction = request.form.get('sanctions')
        students = request.form.get('student')
        complainant = request.form.get('student2')
        date2 = request.form.get('date2')
        current_datetime = datetime.now()
        random_code = generate_random_code()
        current_date = current_datetime.date()
        formatted_date = current_date.strftime("/%m/%d/%Y")

        username = session.get('namestudent', '')

        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor1 = cursor1.cursor()
        db_cursor1.execute(
            "SELECT Name FROM accounts_cics WHERE Username = %s;", (students,))
        result_cics = db_cursor1.fetchone()

        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor1 = cursor1.cursor()
        db_cursor1.execute(
            "SELECT Name FROM accounts_cafad WHERE Username = %s;", (students,))
        result_cafad = db_cursor1.fetchone()

        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor1 = cursor1.cursor()
        db_cursor1.execute(
            "SELECT Name FROM accounts_coe WHERE Username = %s;", (students,))
        result_coe = db_cursor1.fetchone()

        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor1 = cursor1.cursor()
        db_cursor1.execute(
            "SELECT Name FROM accounts_cit WHERE Username = %s;", (students,))
        result_cit = db_cursor1.fetchone()

        if result_cics:
            srcode = students
            name = result_cics[0]

        elif result_cafad:
            srcode = students
            name = result_cafad[0]

        elif result_coe:
            srcode = result_coe[0]
            name = result_coe[0]

        elif result_cit:
            srcode = students
            name = result_cit[0]

        print(srcode+"lol")

        notifs(srcode, "You have a new sanction")

        sanction_mapping = [
            "12.1.1 - attendance, punctuality, cutting classes",
            "12.1.2 - dress code, uniform",
            "12.1.3 - property misuse",
            "12.1.4 - noise disturbance",
            "12.1.5 - posting violation",
            "12.1.6 - notice removal",
            "12.1.7 - littering",
            "12.1.8 - smoking violation",
            "12.1.9 - trespassing",
            "12.1.10 - misconduct",
            "12.1.11 - harassment",
            "12.1.12 - provocation, fight",
            "12.1.13 - PDA",
            "12.1.14 - truancy",
        ]
        sanction_mapping1 = [
            "13.1 - repeat offenses",
            "13.2 - insubordination",
            "13.3 - smoking violation",
            "13.4 - alcohol violation",
            "13.5 - intoxication",
            "13.6 - trespassing",
            "13.7 - property misuse",
            "13.13 - abusive behavior",
            "13.14 - unauthorized membership",
            "13.15 - online misconduct",
            "13.16 - vandalism",
            "13.17 - academic disruption",
            "13.18 - solicitation",
            "13.19 - physical harm",
            "13.20 - weapons possession",
            "13.21 - theft",
            "13.22 - bribery",
            "13.23 - sexual misconduct",
            "13.24 - obscenity",
            "13.25 - defamation",
            "13.26 - physical harm",
            "13.27 - falsification",
            "13.28 - disrepute",
            "13.29 - riot",
            "13.30 - destruction of property",
            "13.31 - burglary",
            "13.32 - hazing",
            "13.33 - drugs",
            "13.34 - firearms possession",
            "13.35 - threats",
            "13.36 - felonies",
            "13.37 - moral turpitude",
        ]
        sanction_mapping2 = [
            "14.1 - cheating, mobile phone",
            "14.2 - cheating, talking",
            "14.3 - cheating, dictating answers",
            "14.4 - cheating, notes possession",
            "14.5 - cheating, outside information",
            "14.6 - cheating, leakage facilitation",
            "14.7 - cheating, buying/selling questions",
            "14.8 - cheating, copying answers",
            "14.9 - cheating, covert devices",
            "14.10 - cheating, impersonation",
            "14.11 - plagiarism",
            "14.12 - cheating, surrogate attendance",
            "14.13 - plagiarism",
            "14.14 - cheating, caught",
            "14.15 - cheating, aiding"
        ]

        sanction_number = None

        if sanction in sanction_mapping:
            sanction_number = "12"
        elif sanction in sanction_mapping1:
            sanction_number = "13"
        elif sanction in sanction_mapping2:
            sanction_number = "14"

        pdf_filename = 'Written Reprimand.docx'
        doc = Document(pdf_filename)

        replace_table_cell_placeholder1(
            doc.tables[0], 2, 8, formatted_date, "(date)")
        replace_table_cell_placeholder1(
            doc.tables[0], 3, 3, name, "(name)")
        replace_table_cell_placeholder1(
            doc.tables[0], 6, 8, sanction_number, "(section)")
        replace_table_cell_placeholder1(doc.tables[0], 7, 2, norms, "norms")
        replace_table_cell_placeholder1(doc.tables[0], 15, 2, username, "NAME")

        doc.save("modified_document.docx")
        pdfpath = os.path.join('modified_document.docx')

        file_name = f'{random_code}_Written Reprimand'
        with open(pdfpath, "rb") as pdf_file:
            pdf_data = pdf_file.read()

        # Insert the report with file information into the database, including file data
        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor = cursor1.cursor()
        db_cursor.execute("INSERT INTO sanctions (sanctions_id,username, course, date_time, sanction, written, written_name,type) VALUES (%s,%s, %s, %s, %s, %s,%s,%s)",
                          (random_code, name, courseorposition, current_datetime, sanction, pdf_data, file_name, kind))
        cursor1.commit()
        db_cursor.close()

        flash('Report submitted successfully')

        return redirect('/head')

    else:
        remarks = request.form.get('remarks')
        norms = request.form.get('norms')
        courseorposition = session.get('course', '')
        print(courseorposition)
        department = request.form.get('department')
        sanction = request.form.get('sanctions')
        students = request.form.get('student')
        effectivity = request.form.get('effectivity')
        checked = request.form.get('checked')
        verified = request.form.get('verified')
        parent = request.form.get('parent')
        days = request.form.get('days')
        current_datetime = datetime.now()
        random_code = generate_random_code()
        current_date = current_datetime.date()
        formatted_date = current_date.strftime("/%m/%d/%Y")

        username = session.get('namestudent', '')

        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor1 = cursor1.cursor()
        db_cursor1.execute(
            "SELECT Name FROM accounts_cics WHERE Username = %s;", (students,))
        result_cics = db_cursor1.fetchone()

        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor1 = cursor1.cursor()
        db_cursor1.execute(
            "SELECT Name FROM accounts_cafad WHERE Username = %s;", (students,))
        result_cafad = db_cursor1.fetchone()

        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor1 = cursor1.cursor()
        db_cursor1.execute(
            "SELECT Name FROM accounts_coe WHERE Username = %s;", (students,))
        result_coe = db_cursor1.fetchone()

        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor1 = cursor1.cursor()
        db_cursor1.execute(
            "SELECT Name FROM accounts_cit WHERE Username = %s;", (students,))
        result_cit = db_cursor1.fetchone()

        if result_cics:
            srcode = students
            name = result_cics[0]

        elif result_cafad:
            srcode = students
            name = result_cafad[0]

        elif result_coe:
            srcode = result_coe[0]
            name = result_coe[0]

        elif result_cit:
            srcode = students
            name = result_cit[0]

        print(srcode+"lol")

        notifs(srcode, "You have a new sanction")

        sanction_mapping = [
            "12.1.1 - attendance, punctuality, cutting classes",
            "12.1.2 - dress code, uniform",
            "12.1.3 - property misuse",
            "12.1.4 - noise disturbance",
            "12.1.5 - posting violation",
            "12.1.6 - notice removal",
            "12.1.7 - littering",
            "12.1.8 - smoking violation",
            "12.1.9 - trespassing",
            "12.1.10 - misconduct",
            "12.1.11 - harassment",
            "12.1.12 - provocation, fight",
            "12.1.13 - PDA",
            "12.1.14 - truancy",
        ]
        sanction_mapping1 = [
            "13.1 - repeat offenses",
            "13.2 - insubordination",
            "13.3 - smoking violation",
            "13.4 - alcohol violation",
            "13.5 - intoxication",
            "13.6 - trespassing",
            "13.7 - property misuse",
            "13.13 - abusive behavior",
            "13.14 - unauthorized membership",
            "13.15 - online misconduct",
            "13.16 - vandalism",
            "13.17 - academic disruption",
            "13.18 - solicitation",
            "13.19 - physical harm",
            "13.20 - weapons possession",
            "13.21 - theft",
            "13.22 - bribery",
            "13.23 - sexual misconduct",
            "13.24 - obscenity",
            "13.25 - defamation",
            "13.26 - physical harm",
            "13.27 - falsification",
            "13.28 - disrepute",
            "13.29 - riot",
            "13.30 - destruction of property",
            "13.31 - burglary",
            "13.32 - hazing",
            "13.33 - drugs",
            "13.34 - firearms possession",
            "13.35 - threats",
            "13.36 - felonies",
            "13.37 - moral turpitude",
        ]
        sanction_mapping2 = [
            "14.1 - cheating, mobile phone",
            "14.2 - cheating, talking",
            "14.3 - cheating, dictating answers",
            "14.4 - cheating, notes possession",
            "14.5 - cheating, outside information",
            "14.6 - cheating, leakage facilitation",
            "14.7 - cheating, buying/selling questions",
            "14.8 - cheating, copying answers",
            "14.9 - cheating, covert devices",
            "14.10 - cheating, impersonation",
            "14.11 - plagiarism",
            "14.12 - cheating, surrogate attendance",
            "14.13 - plagiarism",
            "14.14 - cheating, caught",
            "14.15 - cheating, aiding"
        ]

        sanction_number = None

        if sanction in sanction_mapping:
            sanction_number = "12"
        elif sanction in sanction_mapping1:
            sanction_number = "13"
        elif sanction in sanction_mapping2:
            sanction_number = "14"

        sanction_mapping3 = {
            "12.1.1 - attendance, punctuality, cutting classes": "12.1.1",
            "12.1.2 - dress code, uniform": "12.1.2",
            "12.1.3 - property misuse": "12.1.3",
            "12.1.4 - noise disturbance": "12.1.4",
            "12.1.5 - posting violation": "12.1.5",
            "12.1.6 - notice removal": "12.1.6",
            "12.1.7 - littering": "12.1.7",
            "12.1.8 - smoking violation": "12.1.8",
            "12.1.9 - trespassing": "12.1.9",
            "12.1.10 - misconduct": "12.1.10",
            "12.1.11 - harassment": "12.1.11",
            "12.1.12 - provocation, fight": "12.1.12",
            "12.1.13 - PDA": "12.1.13",
            "12.1.14 - truancy": "12.1.14",
            "13.1 - repeat offenses": "13.1",
            "13.2 - insubordination": "13.2",
            "13.3 - smoking violation": "13.3",
            "13.4 - alcohol violation": "13.4",
            "13.5 - intoxication": "13.5",
            "13.6 - trespassing": "13.6",
            "13.7 - property misuse": "13.7",
            "13.13 - abusive behavior": "13.13",
            "13.14 - unauthorized membership": "13.14",
            "13.15 - online misconduct": "13.15",
            "13.16 - vandalism": "13.16",
            "13.17 - academic disruption": "13.17",
            "13.18 - solicitation": "13.18",
            "13.19 - physical harm": "13.19",
            "13.20 - weapons possession": "13.20",
            "13.21 - theft": "13.21",
            "13.22 - bribery": "13.22",
            "13.23 - sexual misconduct": "13.23",
            "13.24 - obscenity": "13.24",
            "13.25 - defamation": "13.25",
            "13.26 - physical harm": "13.26",
            "13.27 - falsification": "13.27",
            "13.28 - disrepute": "13.28",
            "13.29 - riot": "13.29",
            "13.30 - destruction of property": "13.30",
            "13.31 - burglary": "13.31",
            "13.32 - hazing": "13.32",
            "13.33 - drugs": "13.33",
            "13.34 - firearms possession": "13.34",
            "13.35 - threats": "13.35",
            "13.36 - felonies": "13.36",
            "13.37 - moral turpitude": "13.37",
            "14.1 - cheating, mobile phone": "14.1",
            "14.2 - cheating, talking": "14.2",
            "14.3 - cheating, dictating answers": "14.3",
            "14.4 - cheating, notes possession": "14.4",
            "14.5 - cheating, outside information": "14.5",
            "14.6 - cheating, leakage facilitation": "14.6",
            "14.7 - cheating, buying/selling questions": "14.7",
            "14.8 - cheating, copying answers": "14.8",
            "14.9 - cheating, covert devices": "14.9",
            "14.10 - cheating, impersonation": "14.10",
            "14.11 - plagiarism": "14.11",
            "14.12 - cheating, surrogate attendance": "14.12",
            "14.13 - plagiarism": "14.13",
            "14.14 - cheating, caught": "14.14",
            "14.15 - cheating, aiding": "14.15"

        }
        sanction_number1 = sanction_mapping3.get(sanction, "Unknown")


        pdf_filename = 'letter of suspension.docx'
        doc = Document(pdf_filename)

        replace_table_cell_placeholder1(
            doc.tables[0], 2, 13, formatted_date, "(date)")
        replace_table_cell_placeholder1(
            doc.tables[0], 3, 3, name, "(name)")
        replace_table_cell_placeholder1(
            doc.tables[0], 13, 6, sanction_number1, "(offense)")
        replace_table_cell_placeholder1(doc.tables[0], 13, 14, days, "(days)")
        replace_table_cell_placeholder1(
            doc.tables[0], 14, 4, effectivity, "wew")
        replace_table_cell_placeholder1(
            doc.tables[0], 6, 14, sanction_number, "(section)")
        replace_table_cell_placeholder1(
            doc.tables[0], 18, 3, username, "KRAZY")
        replace_table_cell_placeholder1(
            doc.tables[0], 18, 5, checked, "FERSON")
        replace_table_cell_placeholder1(
            doc.tables[0], 18, 16, verified, "TEST")
        replace_table_cell_placeholder1(
            doc.tables[0], 20, 5, name, "STUDENT")
        replace_table_cell_placeholder1(
            doc.tables[0], 20, 16, verified, "PARENT")

        replace_table_cell_placeholder1(doc.tables[0], 7, 2, norms, "norms")

        doc.save("modified_document.docx")
        pdfpath = os.path.join('modified_document.docx')

        file_name = f'{random_code}_Letter of Suspension'
        with open(pdfpath, "rb") as pdf_file:
            pdf_data = pdf_file.read()

        # Insert the report with file information into the database, including file data
        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor = cursor1.cursor()
        db_cursor.execute("INSERT INTO sanctions (sanctions_id,username, course, date_time, sanction, written, written_name,type) VALUES (%s,%s, %s, %s, %s, %s,%s,%s)",
                          (random_code, name, courseorposition, current_datetime, sanction, pdf_data, file_name, kind))
        cursor1.commit()
        db_cursor.close()

        flash('Report submitted successfully')
        return redirect('head')


@app.route('/submit_approve', methods=['GET', 'POST'])
def submit_approve():
    remarks = request.form.get('remarks')
    report_id = request.form.get('id')
    print(report_id)
    print(remarks)
    status = "Approved"

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()
    db_cursor.execute(
        "UPDATE forms_osd SET remarks = %s, status = %s WHERE form_id = %s", (remarks, status, report_id))
    cursor1.commit()
    db_cursor.close()

    return redirect('/head')


@app.route('/submit_reject', methods=['GET', 'POST'])
def submit_reject():
    remarks = request.form.get('remarks')
    report_id = request.form.get('id')
    print(report_id)
    print(remarks)
    status = "Rejected"

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()
    db_cursor.execute(
        "UPDATE forms_osd SET remarks = %s, status = %s WHERE form_id = %s", (remarks, status, report_id))
    cursor1.commit()
    db_cursor.close()

    return redirect('/head')


@app.route('/delete_sanction', methods=['POST'])
def delete_sanction():
    # Get the sanction ID from the request
    sanction_id = request.form.get('sanctionId')
    print(sanction_id)

    try:
        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor = cursor1.cursor()
        db_cursor.execute(
            "DELETE FROM sanctions WHERE sanctions_id = %s;", (sanction_id,))
        cursor1.commit()
        db_cursor.close()
        return jsonify({"message": "Sanction deleted successfully"})
    except Exception as e:
        # Handle the exception, log the error, and return an error response
        error_message = f"Error deleting sanction: {str(e)}"
        app.logger.error(error_message)
        return jsonify({"error": error_message})


# Make sure to import jsonify from Flask


@app.route('/submit_sanction', methods=['POST'])
def submit_sanction():
    if request.method == 'POST':
        # Get the values from the form
        name = session.get('name')
        course = session.get('course')
        sanction = request.form['sanctions']
        # Get the current date and time
        current_datetime = datetime.now()
        # Get the username from the session

        # Insert the values into the database
        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor = cursor1.cursor()
        db_cursor.execute("INSERT INTO sanctions (username, course, date_time, sanction) VALUES (%s, %s, %s, %s)",
                          (name, course, current_datetime, sanction))
        cursor1.commit()
        db_cursor.close()
        # Optionally, you can redirect to a success page or perform other actions

        return redirect(url_for('homepage'))

    flash('Report submitted successfully!', 'success')
    return redirect(url_for('homepage'))


@app.route('/manage_coord', methods=['GET', 'POST'])
def manage_coord():
    # Retrieve the username from the session if it exists
    username = session.get('username', '')
    user_role = session.get('role', '')
    user_source = session.get('source', '')

    # Query the database to retrieve reports for the logged-in user
    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()

    if user_role == 'accounts_coordinators':
        # If the user is an accounts coordinator, retrieve the course of the user
        db_cursor.execute(
            "SELECT course FROM accounts_coordinators WHERE username = %s", (username,))
        user_course = db_cursor.fetchone()

        if user_course:
            user_course = user_course[0]  # Extract the course from the result

            # Query reports where the course matches the user's course
            db_cursor.execute(
                "SELECT * FROM reports WHERE course = %s", (user_course,))
            reports = db_cursor.fetchall()
    else:
        # For other roles, simply retrieve reports for the logged-in user
        db_cursor.execute(
            "SELECT * FROM reports WHERE username = %s", (username,))
        reports = db_cursor.fetchall()
        user_course = ""

    # Close the cursor
    db_cursor.close()

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor_all = cursor1.cursor()
    db_cursor_all.execute("SELECT * FROM accounts_coordinators")
    coordinators = db_cursor_all.fetchall()
    db_cursor_all.close()

    # Create a dictionary to hold profile pictures as Base64
    profile_pictures = {}

    # Fetch the profile pictures and convert them to Base64 if they exist
    for row in coordinators:
        coord_id = row[0]  # Assuming the first column is the Coord_Id
        image_data = row[3]  # Assuming the fourth column is the image_data

        if image_data:
            profile_picture_base64 = base64.b64encode(
                image_data).decode('utf-8')
            profile_pictures[coord_id] = profile_picture_base64

    return render_template('manage_coord.html', reports=reports, user_source=user_source, user_course=user_course, coordinators=coordinators, profile_pictures=profile_pictures)


def verify_recaptcha(recaptcha_response):
    # Replace with your actual secret key
    secret_key = "6Lf6r8MoAAAAAMqOMUNzyQ--QoeMTyeUcSeBHFCO"

    # Send a POST request to the reCAPTCHA verification endpoint
    response = requests.post(
        "https://www.google.com/recaptcha/api/siteverify",
        {
            "secret": secret_key,
            "response": recaptcha_response,
        },
    )

    result = response.json()

    # Check if reCAPTCHA verification was successful
    if result.get("success"):
        return True
    else:
        return False


tables = [
    'accounts_cics',
    'accounts_cafad',
    'accounts_coe',
    'accounts_cit',
    'accounts_coordinators',
    'accounts_head',
    'accounts_guard'
]


@app.route('/', methods=['GET', 'POST'])
def index():
    # Retrieve the username from the session if it exists
    username = session.get('username', '')

    error_message = None

    if request.method == 'POST':
        # Get the submitted username and password
        submitted_username = request.form['username']
        submitted_password = request.form['password']
        recaptcha_response = request.form['g-recaptcha-response']

        # Verify reCAPTCHA response
        secret_key = '6Lf6r8MoAAAAAMqOMUNzyQ--QoeMTyeUcSeBHFCO'
        recaptcha_url = 'https://www.google.com/recaptcha/api/siteverify'
        response = requests.post(recaptcha_url, data={
            'secret': secret_key,
            'response': recaptcha_response
        })
        #result = "success"

        result = response.json()

        if result['success']:

        #if result == "success" :

            for table in tables:
                query = "SELECT * FROM {} WHERE username = %s AND password = %s".format(
                    table)
                
                cnx = create_connection_pool()
                cursor1=cnx.get_connection()
                db_cursor = cursor1.cursor()
                
                db_cursor.execute(
                    query, (submitted_username, submitted_password))
                result = db_cursor.fetchone()
                db_cursor.close()

                if result:
                    # User exists in the table, set the role and continue
                    session['username'] = submitted_username
                    session['password'] = submitted_password
                    session['role'] = table
                    if 'head' in table:
                        return redirect(url_for('homepage_head'))

                    elif 'coordinator' in table:
                        return redirect(url_for('homepage_head'))

                    else:
                        return redirect(url_for('homepage'))


        else:
            captcha = "Answer the Captcha"
            username = ""
            return render_template('index.html', username=username, captcha=captcha)

    return render_template('index.html', username=username)



@app.route('/menu')
def menu():

    username = session.get('username', '')
    user_role = session.get('role', '')
    user_source = session.get('source', '')

    # Query the database to retrieve reports for the logged-in user
    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()

    if user_role == 'accounts_coordinators':
        # If the user is an accounts coordinator, retrieve the course of the user
        db_cursor.execute(
            "SELECT course FROM accounts_coordinators WHERE username = %s", (username,))
        user_course = db_cursor.fetchone()

        if user_course:
            user_course = user_course[0]  # Extract the course from the result

            # Query reports where the course matches the user's course
            db_cursor.execute(
                "SELECT * FROM reports WHERE course = %s", (user_course,))
            reports = db_cursor.fetchall()

    elif user_role == 'accounts_head':
        db_cursor.execute("SELECT * FROM reports")
        reports = db_cursor.fetchall()
        user_course = ""

    else:
        # For other roles, simply retrieve reports for the logged-in user
        db_cursor.execute(
            "SELECT * FROM reports WHERE username = %s", (username,))
        reports = db_cursor.fetchall()
        user_course = ""

    # Close the cursor
    db_cursor.close()

    return render_template('menu.html', reports=reports, user_source=user_source, user_course=user_course)


@app.route('/request', methods=['GET', 'POST'])
def requestpage():
    # Retrieve the username and role from the session
    username = session.get('username', '')
    user_role = session.get('role', '')
    user_source = session.get('source', '')

    # Query the database to retrieve reports for the logged-in user
    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()

    if user_role == 'accounts_coordinators':
        # If the user is an accounts coordinator, retrieve the course of the user
        db_cursor.execute(
            "SELECT course FROM accounts_coordinators WHERE username = %s", (username,))
        user_course = db_cursor.fetchone()

        if user_course:
            user_course = user_course[0]  # Extract the course from the result

            # Query reports where the course matches the user's course
            db_cursor.execute(
                "SELECT * FROM forms_osd WHERE course = %s", (user_course,))
            reports = db_cursor.fetchall()
    else:
        # For other roles, simply retrieve reports for the logged-in user
        db_cursor.execute(
            "SELECT * FROM forms_osd WHERE username = %s", (username,))
        reports = db_cursor.fetchall()
        user_course = ""

    # Close the cursor
    db_cursor.close()

    return render_template('request.html', reports=reports, user_source=user_source, user_course=user_course)


def is_english(text):
    detectlanguage.configuration.api_key = "9ec41ced9e3687060cbe89995e2b3d51"

    try:

        language_code = detectlanguage.simple_detect(text)

        print(language_code)
        return language_code
    except:
        return False


@app.route('/algorithm/<complaint_text>', methods=['GET', 'POST'])
def algorithm(complaint_text):

    if is_english(complaint_text) == 'en':

        df = pd.read_csv("Data_Gen.csv")

        # Create category_id column
        df['category_id'] = df['offense_tag'].factorize()[0]

        # Text preprocessing and feature extraction
        tfidf = TfidfVectorizer(
            sublinear_tf=True, min_df=5, ngram_range=(1, 2), stop_words='english')
        features = tfidf.fit_transform(df.grievance).toarray()
        labels = df.category_id

        # Train and evaluate the model
        X = df['grievance']
        y = df['offense_tag']

        # Split the data into train and test sets
        X_train, X_test, y_train, y_test = train_test_split(
            X, y, test_size=0.3, random_state=42)

        models = [
            LinearSVC(dual=False),
        ]

        # 5 Cross-validation
        CV = 5
        entries = []
        for model in models:
            model_name = model.__class__.__name__
            accuracies = cross_val_score(
                model, features, labels, scoring='accuracy', cv=CV)
            for fold_idx, accuracy in enumerate(accuracies):
                entries.append((model_name, fold_idx, accuracy))

        model = LinearSVC()
        model.fit(tfidf.transform(X_train), y_train)

        # Sample complaint text
        complaint = complaint_text

        # Predict offenses for the complaint text
        decision_scores = model.decision_function(tfidf.transform([complaint]))

        # Convert decision scores to probabilities using softmax
        def softmax(x):
            exp_x = np.exp(x - np.max(x))
            return exp_x / exp_x.sum(axis=1, keepdims=True)

        predicted_probabilities = softmax(decision_scores)

        # Get the top 10 predicted offenses' category IDs in descending order of prediction score
        top_10_offense_indices = np.argsort(-predicted_probabilities)[0, :5]
        top_10_offense_ids = model.classes_[top_10_offense_indices]

        # Create a dictionary to store the top 10 predicted offenses and their scores
        top_10_offense_scores = {}
        for offense_id, probability in zip(top_10_offense_ids, predicted_probabilities[0, top_10_offense_indices]):
            top_10_offense_scores[offense_id] = round(
                probability * 100)  # Convert to whole number percentage

        # Calculate the total score for the top predicted offenses
        total_score = sum(top_10_offense_scores.values())

        # If the total score is less than 100, distribute the remaining score proportionally
        remaining_score = 100 - total_score
        if remaining_score > 0:
            # Calculate the proportion for each offense based on its probability
            proportions = [probability for _, probability in zip(
                top_10_offense_ids, predicted_probabilities[0, top_10_offense_indices])]
            proportion_sum = sum(proportions)

            # Adjust the scores based on proportions
            for i, offense_id in enumerate(top_10_offense_scores):
                additional_score = round(
                    proportions[i] / proportion_sum * remaining_score)
                top_10_offense_scores[offense_id] += additional_score
                remaining_score -= additional_score
                if remaining_score == 0:
                    break

        print(complaint)
        top_10_offense_scores_list = []
        for offense_id, score in top_10_offense_scores.items():
            print(f"Offense ID: {offense_id}, Score: {score}%")
            top_10_offense_scores_list.append({
                'offense_id': offense_id,
                'score': score,

            })

        type = "english"

        return jsonify(top_10_offense_scores=top_10_offense_scores_list, complaints=complaint, type=type)

    else:
        message = "The report is gibberish or not in English Language"

        type = "gibberish"

        return jsonify(message=message, complaints=complaint_text, type=type)


@app.route('/search_students', methods=['POST'])
def search_students():
    if request.method == 'POST':
        # Updated to match the input name
        search_value = request.form['username']
        session['search_value'] = search_value

        # Perform a database query to search for students in the accounts_cics table
        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor = cursor1.cursor()
        db_cursor.execute(
            "SELECT * FROM accounts_cics WHERE Name LIKE %s", ('%' + search_value + '%',))
        search_results = db_cursor.fetchall()
        db_cursor.close()

        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor = cursor1.cursor()
        db_cursor.execute(
            "SELECT * FROM sanctions WHERE username LIKE %s", ('%' + search_value + '%',))
        search_results = db_cursor.fetchall()
        db_cursor.close()

        # Check if any results were found
        if search_results:
            # Retrieve the first result (you can modify this logic as needed)
            first_result = search_results[0]

            # Extract the name and course from the result
            name = first_result['Name']
            course = first_result['CourseOrPosition']

            # Return the name and course as JSON
            return jsonify({'name': name, 'course': course, 'search_value': session.get('search_value')})
        else:
            # No results found, return an error message as JSON
            print(session.get('name'))
            print(session.get('course'))
            return jsonify({'error': 'No results found'})


@app.route('/forms')
def forms():
    username = session.get('username', '')
    user_role = session.get('role', '')
    user_source = session.get('source', '')
    print(f"Student Name: {user_role}")

    # Query the database to retrieve reports for the logged-in user
    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()

    if user_role == 'accounts_coordinators':
        # If the user is an accounts coordinator, retrieve the course of the user
        db_cursor.execute(
            "SELECT course FROM accounts_coordinators WHERE username = %s", (username,))
        user_course = db_cursor.fetchone()

        if user_course:
            user_course = user_course[0]  # Extract the course from the result

            # Query reports where the course matches the user's course
            db_cursor.execute(
                "SELECT * FROM forms_osd WHERE course = %s", (user_course,))
            reports = db_cursor.fetchall()
    else:
        # For other roles, simply retrieve reports for the logged-in user
        db_cursor.execute(
            "SELECT * FROM forms_osd WHERE username = %s", (username,))
        reports = db_cursor.fetchall()
        user_course = ""

    # Close the cursor
    db_cursor.close()

    return render_template('forms.html', reports=reports, user_source=user_source, user_course=user_course)


@app.route('/download_form/<int:form_id>')
def download_form(form_id):
    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()

    # Retrieve the file data for the given form_id from your database
    db_cursor.execute(
        "SELECT filename, file_data FROM files WHERE id = %s", (form_id,))
    result = db_cursor.fetchone()

    if result is not None:
        filename, file_data = result

        # Serve the file as a downloadable attachment
        response = send_file(
            io.BytesIO(file_data),
            as_attachment=True,
            mimetype='application/pdf',
            download_name=filename + '.pdf'
        )

        return response

    return "Form not found", 404


@app.route('/download_handbook')
def download_handbook():
    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()

    # Retrieve the file data for the given form_id from your database
    db_cursor.execute(
        "SELECT filename, file_data FROM files WHERE id = %s", ("11",))
    result = db_cursor.fetchone()

    if result is not None:
        filename, file_data = result

        # Serve the file as a downloadable attachment
        response = send_file(
            io.BytesIO(file_data),
            as_attachment=True,
            mimetype='application/pdf',
            download_name=filename + '.pdf'
        )

        return response

    return "Form not found", 404


@app.route('/download_manual')
def download_manual():
    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()

    # Retrieve the file data for the given form_id from your database
    db_cursor.execute(
        "SELECT filename, file_data FROM files WHERE id = %s", ("12",))
    result = db_cursor.fetchone()

    if result is not None:
        filename, file_data = result

        # Serve the file as a downloadable attachment
        response = send_file(
            io.BytesIO(file_data),
            as_attachment=True,
            mimetype='application/pdf',
            download_name=filename + '.pdf'
        )

        return response

    return "Form not found", 404


@app.route('/sanctions', methods=['GET', 'POST'])
def sanctions():
    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()
    db_cursor.execute(
        "SELECT * FROM sanctions WHERE username = %s", ("Aedrian Jeao De Torres",))
    sanctions = db_cursor.fetchall()
    print(sanctions)
    db_cursor.close()

    return render_template('homepage.html', sanctions=sanctions)


@app.route('/head', methods=['GET', 'POST'])
def homepage_head():
    success = request.args.get('success')
    username = session.get('username', '')
    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor_reports = cursor1.cursor()

    db_cursor_reports.execute(
            "SELECT * FROM reports WHERE username = %s", (username,))
    complaints1 = db_cursor_reports.fetchall()

    db_cursor_reports.close()
    
    if request.method == 'POST':
        # Handle the POST request (form submission)
        username = request.form['username']
        # Save the username in the session
        session['username'] = username

    # Determine the user source (accounts_cics or accounts_coordinators) and set the user_source variable

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()

    db_cursor.execute(
        "SELECT * FROM accounts_head WHERE username = %s", (username,))
    result_head = db_cursor.fetchone()

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()

    db_cursor.execute(
        "SELECT * FROM accounts_coordinators WHERE username = %s", (username,))
    result_coordinator = db_cursor.fetchone()

    if result_head:
        user_source = 'accounts_head'
        print(user_source)
        session['source'] = user_source

    elif result_coordinator:
        user_source = 'accounts_coordinator'
        print(user_source)
        session['source'] = user_source
    else:
        user_source = 'unknown'  # Handle the case where the user source is not found

    # Close the cursor
    db_cursor.close()

    # Retrieve the profile picture path, name, and course for the logged-in user from the database
    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor1 = cursor1.cursor()

    if user_source == 'accounts_head':
        db_cursor1.execute(
            "SELECT image_data, Name, Position FROM accounts_head WHERE username = %s", (username,))
        role = "head"

    else:
        # Handle the case where user_source is unknown
        db_cursor1.execute(
            "SELECT image_data, name, course FROM accounts_coordinators WHERE username = %s", (username,))
        role = "coordinator"

    result_user_data = db_cursor1.fetchone()

    if role == "head":
        profile_picture_data, name, course = result_user_data
        year = ""
        session['namestudent'] = name
        print(name)

    elif role == "coordinator":
        profile_picture_data, name, course = result_user_data

        year = ""
        session['namestudent'] = name
        session['courseall'] = course
        print(course)
        print(name)

    else:
        # Handle the case where user data is not found
        profile_picture_data = None
        name = "Name not found"
        course = "Course/Position not found"

     # Retrieve the sanctions data within the homepage route
    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor_sanctions = cursor1.cursor()
    db_cursor_sanctions.execute(
        "SELECT * FROM sanctions WHERE username = %s AND type = %s", (name, 'Written Warning'))
    warning = db_cursor_sanctions.fetchall()
    db_cursor_sanctions.close()

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor_sanctions1 = cursor1.cursor()
    db_cursor_sanctions1.execute(
        "SELECT * FROM sanctions WHERE username = %s AND type= %s", (name, "Written Reprimand"))
    reprimand = db_cursor_sanctions1.fetchall()
    db_cursor_sanctions1.close()

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor_sanctions2 = cursor1.cursor()
    db_cursor_sanctions2.execute(
        "SELECT * FROM sanctions WHERE username = %s AND type = %s", (name, "Letter of Suspension"))
    suspension = db_cursor_sanctions2.fetchall()
    db_cursor_sanctions2.close()

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor_call = cursor1.cursor()
    db_cursor_call.execute("SELECT * FROM callslip WHERE coord = %s", (name,))
    call = db_cursor_call.fetchall()
    db_cursor_call.close()


    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor_notice_student = cursor1.cursor()
    db_cursor_notice_student.execute(
        "SELECT * FROM notice_case WHERE name = %s", (name,))
    reports1 = db_cursor_notice_student.fetchall()
    db_cursor_notice_student.close()

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor_notice_complain = cursor1.cursor()
    db_cursor_notice_complain.execute(
        "SELECT * FROM notice_case WHERE complainant = %s", (name,))
    reports2 = db_cursor_notice_complain.fetchall()
    db_cursor_notice_complain.close()

    # Encode the profile picture data as a Base64 string
    if profile_picture_data is not None:
        profile_picture_base643 = base64.b64encode(
            profile_picture_data).decode('utf-8')
    else:
        profile_picture_base64 = None

    username = session.get('username', '')
    user_role = session.get('role', '')
    user_source = session.get('source', '')

    # Query the database to retrieve reports for the logged-in user
    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor_get = cursor1.cursor()
    db_cursor_get1 = cursor1.cursor()

    if user_role == 'accounts_coordinators':
        isCoordinator = "yes"
        db_cursor_get.execute(
            "SELECT course FROM accounts_coordinators WHERE username = %s", (username,))
        user_course = db_cursor_get.fetchone()

        if user_course:
            user_course = user_course[0]  # Extract the course from the result

            # Query reports where the course matches the user's course
            db_cursor_get.execute(
                "SELECT * FROM reports WHERE course = %s", (user_course,))
            reports3 = db_cursor_get.fetchall()
            db_cursor_get1.execute(
                "SELECT * FROM forms_osd WHERE course = %s", (user_course,))
            reports4 = db_cursor_get1.fetchall()

    elif user_role == 'accounts_head':
        isCoordinator = "no"
        db_cursor_get.execute("SELECT * FROM reports")
        reports3 = db_cursor_get.fetchall()
        db_cursor_get1.execute("SELECT * FROM forms_osd")
        reports4 = db_cursor_get1.fetchall()
        user_course = ""

    else:
        # For other roles, simply retrieve reports for the logged-in user
        db_cursor_get.execute(
            "SELECT * FROM reports WHERE username = %s", (username,))
        reports3 = db_cursor_get.fetchall()
        user_course = ""

    # Close the cursor
    db_cursor_get.close()
    db_cursor_get1.close()

    # Create a dictionary to hold profile pictures as Base64

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor_all_cics = cursor1.cursor()
    db_cursor_all_cics.execute("SELECT * FROM accounts_cics")
    cics = db_cursor_all_cics.fetchall()
    db_cursor_all_cics.close()

    profile_pictures1 = {}

    # Fetch the profile pictures and convert them to Base64 if they exist
    for row in cics:
        student_id = row[0]  # Assuming the first column is the Coord_Id
        image_data = row[3]  # Assuming the fourth column is the image_data

        if image_data:
            profile_picture_base64 = base64.b64encode(
                image_data).decode('utf-8')
            profile_pictures1[student_id] = profile_picture_base64

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor_all_cafad = cursor1.cursor()
    db_cursor_all_cafad.execute("SELECT * FROM accounts_cafad")
    cafad = db_cursor_all_cafad.fetchall()
    db_cursor_all_cafad.close()

    profile_pictures2 = {}

    # Fetch the profile pictures and convert them to Base64 if they exist
    for row in cafad:
        student_id = row[0]  # Assuming the first column is the Coord_Id
        image_data = row[3]  # Assuming the fourth column is the image_data

        if image_data:
            profile_picture_base64 = base64.b64encode(
                image_data).decode('utf-8')
            profile_pictures2[student_id] = profile_picture_base64

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor_all_coe = cursor1.cursor()
    db_cursor_all_coe.execute("SELECT * FROM accounts_coe")
    coe = db_cursor_all_coe.fetchall()
    db_cursor_all_coe.close()

    profile_pictures3 = {}

    # Fetch the profile pictures and convert them to Base64 if they exist
    for row in coe:
        student_id = row[0]  # Assuming the first column is the Coord_Id
        image_data = row[3]  # Assuming the fourth column is the image_data

        if image_data:
            profile_picture_base64 = base64.b64encode(
                image_data).decode('utf-8')
            profile_pictures3[student_id] = profile_picture_base64

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor_all_cit = cursor1.cursor()
    db_cursor_all_cit.execute("SELECT * FROM accounts_cit")
    cit = db_cursor_all_cit.fetchall()
    db_cursor_all_cit.close()

    profile_pictures4 = {}

    # Fetch the profile pictures and convert them to Base64 if they exist
    for row in cit:
        student_id = row[0]  # Assuming the first column is the Coord_Id
        image_data = row[3]  # Assuming the fourth column is the image_data

        if image_data:
            profile_picture_base64 = base64.b64encode(
                image_data).decode('utf-8')
            profile_pictures4[student_id] = profile_picture_base64

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor_all_coords = cursor1.cursor()
    db_cursor_all_coords.execute("SELECT * FROM accounts_coordinators")
    coordinators = db_cursor_all_coords.fetchall()
    db_cursor_all_coords.close()

    # Create a dictionary to hold profile pictures as Base64
    profile_pictures = {}

    # Fetch the profile pictures and convert them to Base64 if they exist
    for row in coordinators:
        coord_id = row[0]  # Assuming the first column is the Coord_Id
        image_data = row[3]  # Assuming the fourth column is the image_data

        if image_data:
            profile_picture_base64 = base64.b64encode(
                image_data).decode('utf-8')
            profile_pictures[coord_id] = profile_picture_base64

    username = session.get('courseall', '')
    statuses = ["", "Pending", "On Going", "Rejected", "Case Closed"]
    statuses1 = ["", "Pending", "Ongoing", "Rejected", "Approved"]
    minor = ["12.1.1 - attendance, punctuality, cutting classes", 
             "12.1.2 - dress code, uniform", "12.1.3 - property misuse", 
             "12.1.4 - noise disturbance", "12.1.5 - posting violation",
             "12.1.6 - notice removal","12.1.7 - littering",
             "12.1.8 - smoking violation","12.1.9 - trespassing",
             "12.1.10 - misconduct","12.1.11 - harassment",
             "12.1.12 - provocation, fight","12.1.13 - PDA",
             "12.1.14 - truancy"
             ]
    
    major = ["13.1 - repeat offenses","13.2 - insubordination","13.3 - smoking violation",
             "13.4 - alcohol violation","13.5 - intoxication","13.6 - trespassing",
             "13.7 - property misuse","13.8 - Reckless endangerment","13.9 - Gambling","13.10 - Identity fraud",
             "13.11 - Misuse of university name/logo","13.12 - Unauthorized representation","13.13 - abusive behavior","13.14 - unauthorized membership",
             "13.15 - online misconduct","13.16 - vandalism","13.17 - academic disruption",
             "13.18 - solicitation","13.19 - physical harm","13.20 - weapons possession",
             "13.21 - theft","13.22 - bribery","13.23 - sexual misconduct","13.24 - obscenity",
             "13.25 - defamation","13.26 - physical harm","13.27 - falsification","13.28 - disrepute",
             "13.29 - riot","13.30 - destruction of property","13.31 - burglary","13.32 - hazing",
             "13.33 - drugs","13.34 - firearms possession","13.35 - threats","13.36 - felonies","13.37 - moral turpitude",
             "14.1 - cheating, mobile phone","14.2 - cheating, talking","14.3 - cheating, dictating answers",
             "14.4 - cheating, notes possession","14.5 - cheating, outside information","14.6 - cheating, leakage facilitation",
             "14.7 - cheating, buying/selling questions","14.8 - cheating, copying answers","14.9 - cheating, covert devices",
             "14.10 - cheating, impersonation","14.11 - plagiarism","14.12 - cheating, surrogate attendance",
             "14.13 - plagiarism","14.14 - cheating, caught","14.15 - cheating, aiding"
             ]
    
    counts = {}
    counts1 = {}
    counts2 = {}
    counts3 = {}
    counts4 = {}
    counts5 = {}

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()

    for status in statuses:
        if status:
            query = "SELECT COUNT(*) FROM reports WHERE course = %s AND status = %s"
            db_cursor.execute(query, (username, status))
        else:
            query = "SELECT COUNT(*) FROM reports WHERE course = %s"
            db_cursor.execute(query, (username,))
        result = db_cursor.fetchone()
        counts[status] = result[0]

    db_cursor.close()

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor1 = cursor1.cursor()

    for status in statuses1:
        if status:
            query = "SELECT COUNT(*) FROM forms_osd WHERE course = %s AND status = %s"
            db_cursor1.execute(query, (username, status))
        else:
            query = "SELECT COUNT(*) FROM forms_osd WHERE course = %s"
            db_cursor1.execute(query, (username,))
        result1 = db_cursor1.fetchone()
        counts1[status] = result1[0]

    db_cursor1.close()

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor2 = cursor1.cursor()

    for status in statuses:
        if status:
            query = "SELECT COUNT(*) FROM reports WHERE status = %s"
            db_cursor2.execute(query, (status,))
        else:
            query = "SELECT COUNT(*) FROM reports"
            db_cursor2.execute(query,)
        result2 = db_cursor2.fetchone()
        counts2[status] = result2[0]

    db_cursor2.close()

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor3 = cursor1.cursor()

    for status in statuses1:
        if status:
            query = "SELECT COUNT(*) FROM forms_osd WHERE status = %s"
            db_cursor3.execute(query, (status,))
        else:
            query = "SELECT COUNT(*) FROM forms_osd"
            db_cursor3.execute(query,)
        result3 = db_cursor3.fetchone()
        counts3[status] = result3[0]

    db_cursor3.close()

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor4 = cursor1.cursor()

    for status in minor:
  
        query = "SELECT COUNT(*) FROM sanctions WHERE sanction = %s"
        db_cursor4.execute(query, (status,))

        result4 = db_cursor4.fetchone()
        counts4[status] = result4[0]

    db_cursor4.close()

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor5 = cursor1.cursor()

    for status in major:
      
        query = "SELECT COUNT(*) FROM sanctions WHERE sanction = %s"
        db_cursor5.execute(query, (status,))

        result5 = db_cursor5.fetchone()
        counts5[status] = result5[0]

    db_cursor5.close()

    counts = {
        "Total Number of Complaints": counts[""],
        "Total Pending of Complaints": counts["Pending"],
        "Total On-Going of Complaints": counts["On Going"],
        "Total Rejected of Complaints": counts["Rejected"],
        "Total Resolved of Complaints": counts["Case Closed"]
    }

    counts1 = {
        "Total Number of Requests": counts1[""],
        "Total Pending of Requests": counts1["Pending"],
        "Total On-Going of Requests": counts1["Ongoing"],
        "Total Rejected of Requests": counts1["Rejected"],
        "Total Resolved of Requests": counts1["Approved"]
    }

    counts2 = {
        "Total Number of Complaints": counts2[""],
        "Total Pending of Complaints": counts2["Pending"],
        "Total On-Going of Complaints": counts2["On Going"],
        "Total Rejected of Complaints": counts2["Rejected"],
        "Total Resolved of Complaints": counts2["Case Closed"]
    }

    counts3 = {
        "Total Number of Requests": counts3[""],
        "Total Pending of Requests": counts3["Pending"],
        "Total On-Going of Requests": counts3["Ongoing"],
        "Total Rejected of Requests": counts3["Rejected"],
        "Total Resolved of Requests": counts3["Approved"]
    }

    minor_offenses_sum = 0

    for offense in minor:
        key = f"Total Number of {offense}"
        counts4[key] = counts4.get(offense, 0)
        minor_offenses_sum += counts4[key]

    counts4["Total Number of Minor Offenses"] = minor_offenses_sum

    filtered_counts4 = {key: value for key, value in counts4.items() if "Total Number of" in key}

    major_offenses_sum = 0

    for offense in major:
        key = f"Total Number of {offense}"
        counts5[key] = counts5.get(offense, 0)
        major_offenses_sum += counts5[key]

    counts5["Total Number of Major Offenses"] = major_offenses_sum

    filtered_counts5 = {key: value for key, value in counts5.items() if "Total Number of" in key}


    
    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor_call_student = cursor1.cursor()
    db_cursor_call_student.execute(
        "SELECT * FROM callslip WHERE coord = %s", (name,))
    reports = db_cursor_call_student.fetchall()
    db_cursor_call_student.close()

    
    




    # Pass the sorted offenses, username, profile picture (Base64), name, course, and user_source to the template
    return render_template('homepage_head.html',complaints1=complaints1,counts5=filtered_counts5,counts4= filtered_counts4, counts2=counts2, counts3=counts3, counts=counts, counts1=counts1, isCoordinator=isCoordinator, request=reports4, profile_pictures=profile_pictures, coordinators=coordinators, reports3=reports3, reports1=reports1, reports=reports, reports2=reports2, username=username, profile_picture_base64=profile_picture_base643, name=name, course=course, year=year, user_source=user_source, warning=warning, reprimand=reprimand, suspension=suspension, call=call, profile_pictures1=profile_pictures1, profile_pictures2=profile_pictures2, profile_pictures3=profile_pictures3, profile_pictures4=profile_pictures4, cics=cics, cafad=cafad, coe=coe, cit=cit)


@app.route('/hello', methods=['GET', 'POST'])
def homepage():
    username = session.get('username', '')
    print(username)

    if request.method == 'POST':
        # Handle the POST request (form submission)
        username = request.form['username']
        # Save the username in the session
        session['username'] = username

    # Determine the user source (accounts_cics or accounts_coordinators) and set the user_source variable

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()
    db_cursor.execute(
        "SELECT * FROM accounts_cics WHERE username = %s", (username,))
    result_cics = db_cursor.fetchone()

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()
    db_cursor.execute(
        "SELECT * FROM accounts_cafad WHERE username = %s", (username,))
    result_cafad = db_cursor.fetchone()

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()
    db_cursor.execute(
        "SELECT * FROM accounts_coe WHERE username = %s", (username,))
    result_coe = db_cursor.fetchone()

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()
    db_cursor.execute(
        "SELECT * FROM accounts_cit WHERE username = %s", (username,))
    result_cit = db_cursor.fetchone()

    db_cursor.execute(
        "SELECT * FROM accounts_coordinators WHERE username = %s", (username,))
    result_coordinators = db_cursor.fetchone()

    db_cursor.execute(
        "SELECT * FROM accounts_head WHERE username = %s", (username,))
    result_head = db_cursor.fetchone()

    if result_cics:

        user_source = 'accounts_cics'
        session['source'] = user_source

    elif result_cafad:
        user_source = 'accounts_cafad'
        session['source'] = user_source

    elif result_coe:
        user_source = 'accounts_coe'
        session['source'] = user_source

    elif result_cit:
        user_source = 'accounts_cit'
        session['source'] = user_source

    elif result_coordinators:
        user_source = 'accounts_coordinators'
        session['source'] = user_source

    elif result_head:
        user_source = 'accounts_head'
        print(user_source)
        session['source'] = user_source
    else:
        user_source = 'unknown'  # Handle the case where the user source is not found

    # Close the cursor
    db_cursor.close()

    # Retrieve the profile picture path, name, and course for the logged-in user from the database
    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor1 = cursor1.cursor()

    if user_source == 'accounts_cics':
        db_cursor1.execute(
            "SELECT image_data, Name, Course, Year,role FROM accounts_cics WHERE username = %s", (username,))
        role = "student"

    elif user_source == 'accounts_cafad':
        db_cursor1.execute(
            "SELECT image_data, Name, Course, Year,role FROM accounts_cafad WHERE username = %s", (username,))
        role = "student"

    elif user_source == 'accounts_coe':
        db_cursor1.execute(
            "SELECT image_data, Name, Course, Year,role FROM accounts_coe WHERE username = %s", (username,))
        role = "student"

    elif user_source == 'accounts_cit':
        db_cursor1.execute(
            "SELECT image_data, Name, Course, Year,role FROM accounts_cit WHERE username = %s", (username,))
        role = "student"

    elif user_source == 'accounts_coordinators':
        db_cursor1.execute(
            "SELECT image_data, Name, Course FROM accounts_coordinators WHERE username = %s", (username,))
        role = "coord"

    elif user_source == 'accounts_head':
        db_cursor1.execute(
            "SELECT image_data, Name, Position FROM accounts_head WHERE username = %s", (username,))
        role = "head"

    else:
        # Handle the case where user_source is unknown
        db_cursor1.execute(
            "SELECT image_data, Name FROM accounts_guard WHERE username = %s", (username,))
        role="guard"

    result_user_data = db_cursor1.fetchone()

    if role == "student":
        profile_picture_data, name, course, year, roles = result_user_data
        session['namestudent'] = name
        session['course1'] = course
        print(roles)
        print(name)

    elif role == "coord":
        profile_picture_data, name, course = result_user_data

        year = ""
        session['namestudent'] = name
        session['courseall'] = course
        print(name)

    elif role == "head":
        profile_picture_data, name, course = result_user_data
        year = ""
        session['namestudent'] = name
        print(name)

    elif role == "guard":
        profile_picture_data, name = result_user_data
        course = ""
        year = ""
        roles = "guard"
        print(roles)
        session['namestudent'] = name
        print(name)

    else:
        # Handle the case where user data is not found
        profile_picture_data = None
        name = "Name not found"
        course = "Course/Position not found"

     # Retrieve the sanctions data within the homepage route
    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor_sanctions = cursor1.cursor()
    db_cursor_sanctions.execute(
        "SELECT * FROM sanctions WHERE username = %s AND type = %s", (name, 'Written Warning'))
    warning = db_cursor_sanctions.fetchall()
    db_cursor_sanctions.close()

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor_sanctions1 = cursor1.cursor()
    db_cursor_sanctions1.execute(
        "SELECT * FROM sanctions WHERE username = %s AND type= %s", (name, "Written Reprimand"))
    reprimand = db_cursor_sanctions1.fetchall()
    db_cursor_sanctions1.close()

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor_sanctions2 = cursor1.cursor()
    db_cursor_sanctions2.execute(
        "SELECT * FROM sanctions WHERE username = %s AND type = %s", (name, "Letter of Suspension"))
    suspension = db_cursor_sanctions2.fetchall()
    db_cursor_sanctions2.close()

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor_sanctions3 = cursor1.cursor()
    db_cursor_sanctions3.execute(
        "SELECT * FROM sanctions WHERE username = %s ", (name,))
    sanctions = db_cursor_sanctions3.fetchall()
    db_cursor_sanctions3.close()

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor_call= cursor1.cursor()
    db_cursor_call.execute("SELECT * FROM callslip WHERE coord = %s", (name,))
    call = db_cursor_call.fetchall()
    db_cursor_call.close()

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor_call_student = cursor1.cursor()
    db_cursor_call_student.execute(
        "SELECT * FROM callslip WHERE name = %s", (name,))
    reports = db_cursor_call_student.fetchall()
    db_cursor_call_student.close()

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor_notice_student = cursor1.cursor()
    db_cursor_notice_student.execute(
        "SELECT * FROM notice_case WHERE name = %s", (name,))
    reports1 = db_cursor_notice_student.fetchall()
    db_cursor_notice_student.close()

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor_notice_complain = cursor1.cursor()
    db_cursor_notice_complain.execute(
        "SELECT * FROM notice_case WHERE complainant = %s", (name,))
    reports2 = db_cursor_notice_complain.fetchall()
    db_cursor_notice_complain.close()

    # Encode the profile picture data as a Base64 string
    if profile_picture_data is not None:
        profile_picture_base64 = base64.b64encode(
            profile_picture_data).decode('utf-8')
    else:
        # Handle the case where there is no profile picture data
        profile_picture_base64 = None

    username = session.get('username', '')
    user_role = session.get('role', '')
    user_source = session.get('source', '')

    # Query the database to retrieve reports for the logged-in user
    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor_reports = cursor1.cursor()

    if user_role == 'accounts_coordinators':
        # If the user is an accounts coordinator, retrieve the course of the user
        db_cursor_reports.execute(
            "SELECT course FROM accounts_coordinators WHERE username = %s", (username,))
        user_course = db_cursor_reports.fetchone()

        if user_course:
            user_course = user_course[0]  # Extract the course from the result

            # Query reports where the course matches the user's course
            db_cursor_reports.execute(
                "SELECT * FROM reports WHERE course = %s", (user_course,))
            complaints = db_cursor_reports.fetchall()

    elif user_role == 'accounts_head':
        db_cursor_reports.execute("SELECT * FROM reports")
        complaints = db_cursor_reports.fetchall()
        user_course = ""

    else:
        # For other roles, simply retrieve reports for the logged-in user
        db_cursor_reports.execute(
            "SELECT * FROM reports WHERE username = %s", (username,))
        complaints = db_cursor_reports.fetchall()
        user_course = ""

    # Close the cursor
    db_cursor_reports.close()

    username = session.get('username', '')
    user_role = session.get('role', '')
    user_source = session.get('source', '')

    # Query the database to retrieve reports for the logged-in user
    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor_request = cursor1.cursor()

    if user_role == 'accounts_coordinators':
        # If the user is an accounts coordinator, retrieve the course of the user
        db_cursor_request.execute(
            "SELECT course FROM accounts_coordinators WHERE username = %s", (username,))
        user_course = db_cursor_request.fetchone()

        if user_course:
            user_course = user_course[0]  # Extract the course from the result

            # Query reports where the course matches the user's course
            db_cursor_request.execute(
                "SELECT * FROM forms_osd WHERE course = %s", (user_course,))
            request1 = db_cursor_request.fetchall()
    else:
        # For other roles, simply retrieve reports for the logged-in user
        db_cursor_request.execute(
            "SELECT * FROM forms_osd WHERE username = %s", (username,))
        request1 = db_cursor_request.fetchall()
        user_course = ""

    # Close the cursor
    db_cursor_request.close()

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor_notif = cursor1.cursor()
    db_cursor_notif.execute(
        "SELECT * FROM notifications WHERE user_id = %s", (username,))
    results_notif = db_cursor_notif.fetchall()

    notifs = []  # Create an empty list to store your notifications

    for result_notif in results_notif:
        id, user_id, message, created_at, readit = result_notif
        date_string = created_at.strftime("%Y-%m-%d %H:%M:%S")
        date_obj = datetime.strptime(date_string, "%Y-%m-%d %H:%M:%S")

        formatted_date = date_obj.strftime("%Y-%m-%d")
        formatted_time = date_obj.strftime("%I:%M %p")

        notif = (id, user_id, message, formatted_time, formatted_date, readit)
        notifs.append(notif)  # Append each notification to the list

    db_cursor_notif.close()

    if user_source == 'accounts_cics' or 'accounts_cafad' or 'accounts_coe' or 'accounts_cit':
        user_source = 'student'


    print(roles)

    

    # Pass the sorted offenses, username, profile picture (Base64), name, course, and user_source to the template
    return render_template('homepage.html' ,roles=roles, notif=notifs, sanctions=sanctions, request1=request1, complaints=complaints, reports1=reports1, reports=reports, reports2=reports2, username=username, profile_picture_base64=profile_picture_base64, name=name, course=course, year=year, user_source=user_source, warning=warning, reprimand=reprimand, suspension=suspension, call=call,)


tables1 = [
    'accounts_cics',
    'accounts_cafad',
    'accounts_coe',
    'accounts_cit'
]

def lookup_student_info(username):
    try:
        student_name, student_course = None, None  # Initialize the variables
        cnx = create_connection_pool()
        cursor1 = cnx.get_connection()
        db_cursor = cursor1.cursor(dictionary=True)

        for table in tables1:
            query = f"SELECT Name, Course FROM {table} WHERE username = %s"
            db_cursor.execute(query, (username,))
            result = db_cursor.fetchone()

            if result:
                student_name = result['Name']
                student_course = result['Course']
                # If a match is found in any table, break out of the loop
                break

        return student_name, student_course
    except mysql.connector.Error as err:
        # Handle any errors that may occur during the database query
        print(f"Error: {err}")
        return None, None
    finally:
        db_cursor.close()


# Usage example:


@app.route('/lookup_student', methods=['POST'])
def lookup_student():
    # Get the username from the request
    username = request.form.get('username')

    # Call the function to look up the student's name and course
    student_name, student_course = lookup_student_info(username)
    print(f"Student Name: {student_name}")
    print(f"Student Course: {student_course}")

    session['name'] = student_name
    session['course'] = student_course

    # Return the result as JSON
    student_data = {'Name': student_name, 'CourseOrPosition': student_course}
    return jsonify(student_data)


# Usage example:
@app.route('/count', methods=['POST'])
def count():

    username = session.get('courseall', '')

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()
    db_cursor.execute(
        "SELECT COUNT(*) FROM reports WHERE course = %s", (username,))
    result = db_cursor.fetchone()
    db_cursor.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor1 = cursor1.cursor()
    db_cursor1.execute(
        "SELECT COUNT(*) FROM reports WHERE course = %s AND status = %s", (username, "Pending",))
    result1 = db_cursor1.fetchone()
    db_cursor1.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor2 = cursor1.cursor()
    db_cursor2.execute(
        "SELECT COUNT(*) FROM reports WHERE course = %s AND status = %s", (username, "Ongoing"))
    result2 = db_cursor2.fetchone()
    db_cursor2.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor3 = cursor1.cursor()
    db_cursor3.execute(
        "SELECT COUNT(*) FROM reports WHERE course = %s AND status = %s", (username, "Rejected"))
    result3 = db_cursor3.fetchone()
    db_cursor3.close

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor4 = cursor1.cursor()
    db_cursor4.execute(
        "SELECT COUNT(*) FROM reports WHERE course = %s AND status = %s", (username, "Case Closed"))
    result4 = db_cursor4.fetchone()
    db_cursor4.close

    countreports = result[0]
    countpending = result1[0]
    countongoing = result2[0]
    countrejected = result3[0]
    countcaseclosed = result4[0]

    print(countreports)

    # Return the result as JSON
    student_data = {'Reports': countreports, 'Pending': countpending,
                    'Ongoing': countongoing, 'Rejected': countrejected, 'Caseclosed': countcaseclosed}
    return jsonify(student_data)


@app.route('/check', methods=['POST'])
def check():

    username = session.get('namestudent', '')
    print(username)

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()
    db_cursor.execute(
        "SELECT COUNT(*) FROM callslip WHERE name = %s", (username,))
    result = db_cursor.fetchone()

    checks = result[0]


    if checks:
        tf = "true"
        session['oneshow'] = "true"

    else:
        tf = "false"
        session['oneshow'] = "true"

    # Return the result as JSON
    student_data = {'Reports': tf}
    return jsonify(student_data)


@app.route('/check2', methods=['POST'])
def check2():

    lol = "false"

    oneshow = session.get('oneshow', '')

    print(oneshow)

    if oneshow == "true":
        lol = "true"

    else:
        lol = "false"

    # Return the result as JSON
    student_data = {'show': lol}
    return jsonify(student_data)


@app.route('/download_report_file/<string:report_id>')
def download_report_file(report_id):
    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()
    db_cursor.execute(
        "SELECT file_form, file_form_name FROM reports WHERE report_id = %s", (report_id,))
    result = db_cursor.fetchone()

    if result is not None:
        file_data, file_name = result

        # Set the content type header to PDF
        content_type = 'application/pdf'

        # Set the filename to "default.pdf"
        response = make_response(file_data)
        response.headers['Content-Type'] = content_type

        # Set the filename to "default.pdf"
        response.headers['Content-Disposition'] = f'attachment; filename="{file_name}.pdf"'

        # Close the cursor after fetching the result
        db_cursor.close()

        return response

    # Handle the case where the file is not found
    db_cursor.close()
    return "File not found", 404


@app.route('/download_supporting_document/<string:report_id>')
def download_supporting_document(report_id):
    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()
    db_cursor.execute(
        "SELECT file_support_name, file_support_type, file_support FROM reports WHERE report_id = %s", (report_id,))
    result = db_cursor.fetchone()

    if result is not None:
        file_name, file_type, file_data = result

        # Set the content type header based on the supporting document's type stored in the database
        content_type = file_type

        # Set the filename to have the original name and extension
        response = make_response(file_data)
        response.headers['Content-Type'] = content_type

        # Set the filename based on the stored name and extension
        response.headers['Content-Disposition'] = f'attachment; filename="{file_name}{file_type}"'

        # Close the cursor after fetching the result
        db_cursor.close()

        return response

    # Handle the case where the file is not found
    db_cursor.close()
    return "File not found", 404


@app.route('/download_report_file1/<string:report_id>')
def download_report_file1(report_id):
    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()
    db_cursor.execute(
        "SELECT file_form, file_form_name FROM forms_osd WHERE form_id = %s", (report_id,))
    result = db_cursor.fetchone()

    if result is not None:
        file_data, file_name = result

        # Set the content type header to PDF
        content_type = 'application/pdf'

        # Set the filename to "default.pdf"
        response = make_response(file_data)
        response.headers['Content-Type'] = content_type

        # Set the filename to "default.pdf"
        response.headers['Content-Disposition'] = f'attachment; filename="{file_name}.pdf"'

        # Close the cursor after fetching the result
        db_cursor.close()

        return response

    # Handle the case where the file is not found
    db_cursor.close()
    return "File not found", 404


@app.route('/download_supporting_document1/<string:report_id>')
def download_supporting_document1(report_id):
    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()
    db_cursor.execute(
        "SELECT file_support_name, file_support_type, file_support FROM forms_osd WHERE form_id = %s", (report_id,))
    result = db_cursor.fetchone()

    if result is not None:
        file_name, file_type, file_data = result

        # Set the content type header based on the supporting document's type stored in the database
        content_type = file_type

        # Set the filename to have the original name and extension
        response = make_response(file_data)
        response.headers['Content-Type'] = content_type

        # Set the filename based on the stored name and extension
        response.headers['Content-Disposition'] = f'attachment; filename="{file_name}{file_type}"'

        # Close the cursor after fetching the result
        db_cursor.close()

        return response

    # Handle the case where the file is not found
    db_cursor.close()
    return "File not found", 404


@app.route('/change_report_status/<string:report_id>', methods=['POST'])
def change_report_status(report_id):
    new_status = request.form['new_status']
    print(new_status)
    print(report_id)

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor1 = cursor1.cursor()
    db_cursor1.execute(
        "SELECT report_id FROM reports WHERE id = %s;", (report_id,))
    result = db_cursor1.fetchone()
    results = result[0]
    db_cursor1.close()

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor2 = cursor1.cursor()
    db_cursor2.execute(
        "SELECT username FROM reports WHERE id = %s;", (report_id,))
    result1 = db_cursor2.fetchone()
    results1 = result1[0]

    notifs(results1, "Your Report ID."+results+" has change status")

    
    db_cursor2.close()

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()
    db_cursor.execute(
        "UPDATE reports SET status = %s WHERE id = %s;", (new_status, report_id))
    cursor1.commit()  # Make sure to commit the changes to the database
    db_cursor.close()
    db_cursor.close()

    flash('Status has been successfully changed', 'success')

    return redirect(url_for('homepage_head'))


@app.route('/change_report_status1/<string:report_id>', methods=['POST'])
def change_report_status1(report_id):
    new_status = request.form['new_status']
    print(new_status)
    print(report_id)

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor2 = cursor1.cursor()
    db_cursor2.execute(
        "SELECT username FROM forms_osd WHERE form_id = %s;", (report_id,))
    result1 = db_cursor2.fetchone()
    results1 = result1[0]

    notifs(results1, "Your Report ID."+report_id+" has change status")

    db_cursor2.close()

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()
    db_cursor.execute(
        "UPDATE forms_osd SET status = %s WHERE form_id = %s;", (new_status, report_id))
    cursor1.commit()  # Make sure to commit the changes to the database
    db_cursor.close()

    flash('Status has been successfully changed', 'success')

    return redirect(url_for('homepage_head'))


@app.route('/delete_call/<string:report_id>', methods=['POST'])
def delete_call(report_id):
    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()
    db_cursor.execute("DELETE FROM callslip WHERE call_id = %s;", (report_id,))
    cursor1.commit()  # Make sure to commit the changes to the database
    db_cursor.close()

    return redirect(url_for('homepage'))


@app.route('/delete_report/<string:report_id>', methods=['POST'])
def delete_report(report_id):
    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()
    db_cursor.execute(
        "DELETE FROM reports WHERE report_id = %s;", (report_id,))
    cursor1.commit()  # Make sure to commit the changes to the database
    db_cursor.close()

    return redirect(url_for('homepage_head'))


@app.route('/delete_report1/<string:report_id>', methods=['POST'])
def delete_report1(report_id):
    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()
    db_cursor.execute(
        "DELETE FROM forms_osd WHERE form_id = %s;", (report_id,))
    cursor1.commit()  # Make sure to commit the changes to the database
    db_cursor.close()

    return redirect(url_for('homepage_head'))

@app.route('/delete_report2/<string:report_id>', methods=['POST'])
def delete_report2(report_id):
    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()
    db_cursor.execute(
        "DELETE FROM callslip WHERE call_id = %s;", (report_id,))
    cursor1.commit()  # Make sure to commit the changes to the database
    db_cursor.close()

    return redirect(url_for('homepage_head'))


@app.route('/delete_all_report1/<string:status>', methods=['POST'])
def delete_all_report1(status):

    print(status)

    if status == "Result":
        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor = cursor1.cursor()
        db_cursor.execute(
            "DELETE FROM forms_osd WHERE  status = %s", ("Approved",))
        cursor1.commit()  # Make sure to commit the changes to the database
        db_cursor.close()


        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor = cursor1.cursor()
        db_cursor.execute(
            "DELETE FROM forms_osd WHERE  status = %s", ("Rejected",))
        cursor1.commit()  # Make sure to commit the changes to the database
        db_cursor.close()


        return redirect(url_for('homepage_head'))


    else:
        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor = cursor1.cursor()
        db_cursor.execute(
            "DELETE FROM forms_osd WHERE  status = %s;", (status,))
        cursor1.commit()  # Make sure to commit the changes to the database
        db_cursor.close()
        return redirect(url_for('homepage_head'))


    


@app.route('/delete_all_report/<string:report_id>', methods=['POST'])
def delete_all_report(report_id):
    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()
    db_cursor.execute("DELETE FROM reports WHERE course = %s;", (report_id,))
    cursor1.commit()  # Make sure to commit the changes to the database
    db_cursor.close()

    return redirect(url_for('homepage_head'))


@app.route('/delete_all_report2/', methods=['POST'])
def delete_all_report2():
    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()
    db_cursor.execute("DELETE FROM reports")
    cursor1.commit()  # Make sure to commit the changes to the database
    db_cursor.close()

    return redirect(url_for('homepage_head'))


@app.route('/delete-notification', methods=['POST'])
def delete_notification():

    notification_id = request.form.get('id')
    print(notification_id)

    cnx = create_connection_pool()
    cursor1=cnx.get_connection()
    db_cursor = cursor1.cursor()
    db_cursor.execute(
        "DELETE FROM notifications WHERE id = %s", (notification_id,))
    cursor1.commit()
    db_cursor.close()

    return 'Notification deleted successfully'


@app.route('/lookup_sanctions', methods=['POST'])
def lookup_sanctions():
    if request.method == 'POST':
        name = session.get('name', '')  # Updated to match the input name

        # Perform a database query to search for sanctions based on the username
        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor = cursor1.cursor()
        db_cursor.execute(
            "SELECT date_time, sanction, sanctions_id, written_name FROM sanctions WHERE Username LIKE %s", ('%' + name + '%',))
        search_sanctions = db_cursor.fetchall()
        db_cursor.close()

        # Check if any sanctions were found
        if search_sanctions:
            # Convert datetime objects to string representations
            formatted_sanctions = [{'date_time': str(
                entry[0]), 'sanction': entry[1], 'sanctions_id': entry[2], 'written_name': entry[3]} for entry in search_sanctions]
            return jsonify({'sanctions': formatted_sanctions})
        else:
            return jsonify({'error': 'No sanctions found'})


@app.route('/logout', methods=['GET'])
def logout():
    # Clear the session data
    session.clear()

    # Redirect the user to the login page or any other appropriate page
    return redirect('/')


@app.route('/preview_call_file', methods=['GET','POST'])
def preview_call_file():
    password = session.get('password', '')
    print(password)
    password1 = request.form.get('complainant')
    print(password1)
    report_id = request.form.get('id')
    print(report_id)
    db_cursor = None  # Initialize db_cursor to None


    if password == password1:

        print("wow")


        try:
            cnx = create_connection_pool()
            cursor1=cnx.get_connection()
            db_cursor = cursor1.cursor()
            db_cursor.execute("SELECT file, file_name FROM callslip WHERE call_id = %s", (report_id,))
            result = db_cursor.fetchone()

            

            if result:
                print("wow1")
                file_content, file_type = result

                response = send_file(
                    io.BytesIO(file_content),
                    mimetype='application/octet-stream',
                )

                response.headers['Content-Disposition'] = f'inline; filename={file_type}.docx'

                db_cursor.close()

                return response

        except Exception as e:
            # Handle any exceptions, e.g., log the error
            pass  # Add your error handling code here

        # Handle the case where the file was not found
        return "File not found", 404,

    else:

        flash('Invalid password')
        return redirect('/head')


@app.route('/preview_notice_file', methods=['GET','POST'])
def preview_notice_file():
    password = session.get('password', '')
    print(password)
    password1 = request.form.get('complainant')
    print(password1)
    report_id = request.form.get('id')
    print(report_id)
    db_cursor = None  # Initialize db_cursor to None


    if password == password1:

        print("wow")


        try:
            cnx = create_connection_pool()
            cursor1=cnx.get_connection()
            db_cursor = cursor1.cursor()
            db_cursor.execute("SELECT file, file_name FROM notice_case WHERE notice_id = %s", (report_id,))
            result = db_cursor.fetchone()

            

            if result:
                print("wow1")
                file_content, file_type = result

                response = send_file(
                    io.BytesIO(file_content),
                    mimetype='application/octet-stream',
                )

                response.headers['Content-Disposition'] = f'inline; filename={file_type}.docx'

                db_cursor.close()

                return response 

        except Exception as e:
            # Handle any exceptions, e.g., log the error
            pass  # Add your error handling code here

        # Handle the case where the file was not found
        return "File not found", 404,

    else:

        flash('Invalid password')
        return redirect('/hello')


@app.route('/preview_written_file', methods=['GET','POST'])
def preview_written_file():
    password = session.get('password', '')
    print(password)
    password1 = request.form.get('complainant')
    print(password1)
    report_id = request.form.get('id')
    print(report_id)
    db_cursor = None  # Initialize db_cursor to None


    if password == password1:

        print("wow")


        try:
            cnx = create_connection_pool()
            cursor1=cnx.get_connection()
            db_cursor = cursor1.cursor()
            db_cursor.execute("SELECT written, written_name FROM sanctions WHERE sanctions_id = %s", (report_id,))
            result = db_cursor.fetchone()

            

            if result:
                print("wow1")
                file_content, file_type = result

                response = send_file(
                    io.BytesIO(file_content),
                    mimetype='application/octet-stream',
                )

                response.headers['Content-Disposition'] = f'inline; filename={file_type}.docx'

                db_cursor.close()

                return response

        except Exception as e:
            # Handle any exceptions, e.g., log the error
            pass  # Add your error handling code here

        # Handle the case where the file was not found
        return "File not found", 404,

    else:

        flash('Invalid password')
        return redirect('/head')


@app.route('/preview_report_file', methods=['GET','POST'])
def preview_report_file():
    password = session.get('password', '')
    print(password)
    password1 = request.form.get('complainant')
    print(password1)
    report_id = request.form.get('id')
    print(report_id)
    db_cursor = None  # Initialize db_cursor to None


    if password == password1:

        print("wow")


        try:
            cnx = create_connection_pool()
            cursor1=cnx.get_connection()
            db_cursor = cursor1.cursor()
            db_cursor.execute("SELECT file_form, file_form_name FROM reports WHERE id = %s", (report_id,))
            result = db_cursor.fetchone()

            

            if result:
                print("wow1")
                file_content, file_type = result

                

                response = send_file(
                    io.BytesIO(file_content),
                    mimetype='application/octet-stream',
                )

                response.headers['Content-Disposition'] = f'inline; filename={file_type}.docx'

                db_cursor.close()
                

                return response

        except Exception as e:
            # Handle any exceptions, e.g., log the error
            pass  # Add your error handling code here

        # Handle the case where the file was not found
        return "File not found", 404,

    else:

        flash('Invalid password')
        return redirect('/head')
        





@app.route('/preview_support_file/<string:report_id>/<int:idx>', methods=['GET'])
def preview_support_file(report_id,idx):
    db_cursor = None  # Initialize db_cursor to None

    try:
        if idx == 1:
            cnx = create_connection_pool()
            cursor1=cnx.get_connection()
            db_cursor = cursor1.cursor()
            db_cursor.execute(
                "SELECT file_support FROM reports WHERE report_id = %s", (report_id,))
            file_content = db_cursor.fetchone()
            db_cursor.execute(
                "SELECT file_support_name FROM reports WHERE report_id = %s", (report_id,))
            file_type = db_cursor.fetchone()

            file_type = file_type[0]

            if file_content:
                file_content = file_content[0]

                response = send_file(
                    io.BytesIO(file_content),
                    mimetype='application/octet-stream',
                )

                response.headers['Content-Disposition'] = f'inline; filename={file_type}'

                return response
        elif idx == 2:
            cnx = create_connection_pool()
            cursor1=cnx.get_connection()
            db_cursor = cursor1.cursor()
            db_cursor.execute(
                "SELECT file_support1 FROM reports WHERE report_id = %s", (report_id,))
            file_content = db_cursor.fetchone()
            db_cursor.execute(
                "SELECT file_support_name1 FROM reports WHERE report_id = %s", (report_id,))
            file_type = db_cursor.fetchone()

            file_type = file_type[0]

            if file_content:
                file_content = file_content[0]

                response = send_file(
                    io.BytesIO(file_content),
                    mimetype='application/octet-stream',
                )

                response.headers['Content-Disposition'] = f'inline; filename={file_type}'

                return response

        else:
            cnx = create_connection_pool()
            cursor1=cnx.get_connection()
            db_cursor = cursor1.cursor()
            db_cursor.execute(
                "SELECT file_support2 FROM reports WHERE report_id = %s", (report_id,))
            file_content = db_cursor.fetchone()
            db_cursor.execute(
                "SELECT file_support_name2 FROM reports WHERE report_id = %s", (report_id,))
            file_type = db_cursor.fetchone()

            file_type = file_type[0]

            if file_content:
                file_content = file_content[0]

                response = send_file(
                    io.BytesIO(file_content),
                    mimetype='application/octet-stream',
                )

                response.headers['Content-Disposition'] = f'inline; filename={file_type}'

                return response   
        

    except Exception as e:
        # Handle any exceptions, e.g., log the error
        pass  # Add your error handling code here

    finally:
        if db_cursor is not None:
            db_cursor.close()  # Close the cursor if it's not None

    # Handle the case where the file was not found
    return "File not found", 404


@app.route('/preview_report_file1', methods=['GET','POST'])
def preview_report_file1():
    password = session.get('password', '')
    print(password)
    password1 = request.form.get('complainant')
    print(password1)
    report_id = request.form.get('id')
    print(report_id)
    db_cursor = None  # Initialize db_cursor to None


    if password == password1:

        print("wow")


        try:
            cnx = create_connection_pool()
            cursor1=cnx.get_connection()
            db_cursor = cursor1.cursor()
            db_cursor.execute("SELECT file_form, file_form_name FROM forms_osd WHERE form_id = %s", (report_id,))
            result = db_cursor.fetchone()

            

            if result:
                print("wow1")
                file_content, file_type = result

                response = send_file(
                    io.BytesIO(file_content),
                    mimetype='application/octet-stream',
                )

                response.headers['Content-Disposition'] = f'inline; filename={file_type}.docx'

                db_cursor.close()

                return response

        except Exception as e:
            # Handle any exceptions, e.g., log the error
            pass  # Add your error handling code here

        # Handle the case where the file was not found
        return "File not found", 404,

    else:

        flash('Invalid password')
        return redirect('/head')
        


@app.route('/preview_support_file1/<string:report_id>', methods=['GET'])
def preview_support_file1(report_id):
    db_cursor = None  # Initialize db_cursor to None

    try:
        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor = cursor1.cursor()
        db_cursor.execute(
            "SELECT file_support FROM forms_osd WHERE form_id = %s", (report_id,))
        file_content = db_cursor.fetchone()
        db_cursor.execute(
            "SELECT file_support_name FROM forms_osd WHERE form_id = %s", (report_id,))
        file_type = db_cursor.fetchone()

        file_type = file_type[0]

        if file_content:
            file_content = file_content[0]

            response = send_file(
                io.BytesIO(file_content),
                mimetype='application/octet-stream',
            )

            response.headers['Content-Disposition'] = f'inline; filename={file_type}'

            return response

    except Exception as e:
        # Handle any exceptions, e.g., log the error
        pass  # Add your error handling code here

    finally:
        if db_cursor is not None:

            db_cursor.close()  # Close the cursor if it's not None

    # Handle the case where the file was not found
    return "File not found", 404,


@app.route('/update-database', methods=['POST'])
def update_database():
    try:
        # Get the JSON data from the request
        data = request.get_json()

        id = data.get('coordId')
        username = data.get('username')
        password = data.get('password')
        profile_pic_base64 = data.get('picId')
        name = data.get('name')
        course = data.get('course')

        if profile_pic_base64:
            profile_pic = base64.b64decode(profile_pic_base64)
        else:
            profile_pic = None  # Handle the case where there is no profile picture

        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor = cursor1.cursor()
        db_cursor.execute("UPDATE accounts_coordinators SET username = %s, password = %s, image_data = %s, name = %s, course = %s WHERE id = %s;",
                          (username, password, profile_pic, name, course, id))
        cursor1.commit()  # Make sure to commit the changes to the database
        db_cursor.close()

        return jsonify({"message": "Database updated successfully"})
    except Exception as e:
        # Handle any errors that may occur during the update
        return jsonify({"error": str(e)})


@app.route('/update-database1', methods=['POST'])
def update_database1():
    try:
        # Get the JSON data from the request
        data = request.get_json()
        print("data")

        id = data.get('coordId')
        print(id)
        username = data.get('username')
        password = data.get('password')
        profile_pic_base64 = data.get('picId')
        name = data.get('name')
        print(name)
        course = data.get('course')

        if profile_pic_base64:
            profile_pic = base64.b64decode(profile_pic_base64)
        else:
            profile_pic = None  # Handle the case where there is no profile picture

        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor = cursor1.cursor()
        db_cursor.execute("UPDATE accounts_cics SET username = %s, password = %s, image_data = %s, name = %s, course = %s WHERE id = %s;",
                          (username, password, profile_pic, name, course, id))
        cursor1.commit()  # Make sure to commit the changes to the database
        db_cursor.close()

        return jsonify({"message": "Database updated successfully"})
    except Exception as e:
        # Handle any errors that may occur during the update
        return jsonify({"error": str(e)})


@app.route('/update-database2', methods=['POST'])
def update_database2():
    try:
        # Get the JSON data from the request
        data = request.get_json()

        id = data.get('coordId')
        username = data.get('username')
        password = data.get('password')
        profile_pic_base64 = data.get('picId')
        name = data.get('name')
        course = data.get('course')

        if profile_pic_base64:
            profile_pic = base64.b64decode(profile_pic_base64)
        else:
            profile_pic = None  # Handle the case where there is no profile picture

        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor = cursor1.cursor()
        db_cursor.execute("UPDATE accounts_cit SET username = %s, password = %s, image_data = %s, name = %s, course = %s WHERE id = %s;",
                          (username, password, profile_pic, name, course, id))
        cursor1.commit()  # Make sure to commit the changes to the database
        db_cursor.close()

        return jsonify({"message": "Database updated successfully"})
    except Exception as e:
        # Handle any errors that may occur during the update
        return jsonify({"error": str(e)})


@app.route('/update-database3', methods=['POST'])
def update_database3():
    try:
        # Get the JSON data from the request
        data = request.get_json()

        id = data.get('coordId')
        username = data.get('username')
        password = data.get('password')
        profile_pic_base64 = data.get('picId')
        name = data.get('name')
        course = data.get('course')

        if profile_pic_base64:
            profile_pic = base64.b64decode(profile_pic_base64)
        else:
            profile_pic = None  # Handle the case where there is no profile picture

        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor = cursor1.cursor()
        db_cursor.execute("UPDATE accounts_cafad SET username = %s, password = %s, image_data = %s, name = %s, course = %s WHERE id = %s;",
                          (username, password, profile_pic, name, course, id))
        cursor1.commit()  # Make sure to commit the changes to the database
        db_cursor.close()

        return jsonify({"message": "Database updated successfully"})
    except Exception as e:
        # Handle any errors that may occur during the update
        return jsonify({"error": str(e)})


@app.route('/update-database4', methods=['POST'])
def update_database4():
    try:
        # Get the JSON data from the request
        data = request.get_json()

        id = data.get('coordId')
        username = data.get('username')
        password = data.get('password')
        profile_pic_base64 = data.get('picId')
        name = data.get('name')
        course = data.get('course')

        if profile_pic_base64:
            profile_pic = base64.b64decode(profile_pic_base64)
        else:
            profile_pic = None  # Handle the case where there is no profile picture

        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor = cursor1.cursor()
        db_cursor.execute("UPDATE accounts_coe SET username = %s, password = %s, image_data = %s, name = %s, course = %s WHERE id = %s;",
                          (username, password, profile_pic, name, course, id))
        cursor1.commit()  # Make sure to commit the changes to the database
        db_cursor.close()

        return jsonify({"message": "Database updated successfully"})
    except Exception as e:
        # Handle any errors that may occur during the update
        return jsonify({"error": str(e)})


@app.route('/edit_pic', methods=['POST'])
def edit_pic():
    ids = request.form.get('id')
    print(ids)
    try:
        pic = request.files['file3']

        if pic:
            # Read the image data from the file
            image_data = memoryview(pic.read()).tobytes()
        else:
            image_data = None  # Handle the case where there is no profile picture

        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor = cursor1.cursor()
        db_cursor.execute(
            "UPDATE accounts_coordinators SET image_data = %s WHERE id = %s;", (image_data, ids))
        cursor1.commit()  # Make sure to commit the changes to the database
        db_cursor.close()

        return redirect(url_for('homepage_head'))
    except Exception as e:
        # Handle any errors that may occur during the update
        return jsonify({"error": str(e)})


@app.route('/edit_pic1', methods=['POST'])
def edit_pic1():
    ids = request.form.get('id')
    print(ids)
    try:
        pic = request.files['file3']

        if pic:
            # Read the image data from the file
            image_data = memoryview(pic.read()).tobytes()
        else:
            image_data = None  # Handle the case where there is no profile picture

        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor = cursor1.cursor()
        db_cursor.execute(
            "UPDATE accounts_cics SET image_data = %s WHERE id = %s;", (image_data, ids))
        cursor1.commit()  # Make sure to commit the changes to the database
        db_cursor.close()

        return redirect(url_for('homepage_head'))
    except Exception as e:
        # Handle any errors that may occur during the update
        return jsonify({"error": str(e)})


@app.route('/edit_pic2', methods=['POST'])
def edit_pic2():
    ids = request.form.get('id')
    print(ids)
    try:
        pic = request.files['file3']

        if pic:
            # Read the image data from the file
            image_data = memoryview(pic.read()).tobytes()
        else:
            image_data = None  # Handle the case where there is no profile picture

        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor = cursor1.cursor()
        db_cursor.execute(
            "UPDATE accounts_cit SET image_data = %s WHERE id = %s;", (image_data, ids))
        cursor1.commit()  # Make sure to commit the changes to the database
        db_cursor.close()

        return redirect(url_for('homepage_head'))
    except Exception as e:
        # Handle any errors that may occur during the update
        return jsonify({"error": str(e)})
    



@app.route('/edit_pic3', methods=['POST'])
def edit_pic3():
    ids = request.form.get('id')
    print(ids)
    try:
        pic = request.files['file3']

        if pic:
            # Read the image data from the file
            image_data = memoryview(pic.read()).tobytes()
        else:
            image_data = None  # Handle the case where there is no profile picture

        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor = cursor1.cursor()
        db_cursor.execute(
            "UPDATE accounts_cafad SET image_data = %s WHERE id = %s;", (image_data, ids))
        cursor1.commit()  # Make sure to commit the changes to the database
        db_cursor.close()

        return redirect(url_for('homepage_head'))
    except Exception as e:
        # Handle any errors that may occur during the update
        return jsonify({"error": str(e)})
    



@app.route('/edit_pic4', methods=['POST'])
def edit_pic4():
    ids = request.form.get('id')
    print(ids)
    try:
        pic = request.files['file3']

        if pic:
            # Read the image data from the file
            image_data = memoryview(pic.read()).tobytes()
        else:
            image_data = None  # Handle the case where there is no profile picture

        cnx = create_connection_pool()
        cursor1=cnx.get_connection()
        db_cursor = cursor1.cursor()
        db_cursor.execute(
            "UPDATE accounts_coe SET image_data = %s WHERE id = %s;", (image_data, ids))
        cursor1.commit()  # Make sure to commit the changes to the database
        db_cursor.close()

        return redirect(url_for('homepage_head'))
    except Exception as e:
        # Handle any errors that may occur during the update
        return jsonify({"error": str(e)})





if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0')

